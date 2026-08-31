"""Build a bounded delivery report from the completed matrix and exploratory evidence.

This is intentionally separate from the historical Task 17/18 acceptance
report builder.  The customer-authorized path removed the pseudo-label gate;
the report therefore records the supervised reference matrix and the
exploratory Student/open-world evidence without inventing a formal SSOD
acceptance claim.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import tempfile
from pathlib import Path
from typing import Any, Mapping


class DeliveryReportError(ValueError):
    """Raised when required evidence is absent or inconsistent."""


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _load(path: Path, label: str) -> Mapping[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError) as error:
        raise DeliveryReportError(f"{label} cannot be read: {error}") from error
    if not isinstance(value, Mapping):
        raise DeliveryReportError(f"{label} is not a JSON object")
    return value


def _evidence(path: Path, label: str) -> dict[str, Any]:
    path = path.resolve(strict=True)
    if not path.is_file():
        raise DeliveryReportError(f"{label} is missing: {path}")
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": _sha256(path)}


def _metric_text(value: object) -> str:
    return f"{float(value):.6f}" if isinstance(value, (int, float)) and not isinstance(value, bool) else "n/a"


def _validate_matrix(matrix: Mapping[str, Any]) -> list[Mapping[str, Any]]:
    if matrix.get("protocol") != "task12_supervised_reference_matrix":
        raise DeliveryReportError("supervised matrix protocol is unsupported")
    summary = matrix.get("summary")
    if not isinstance(summary, Mapping) or summary.get("complete_runs") != 6 or summary.get("failed_runs") != 0:
        raise DeliveryReportError("supervised matrix is incomplete or contains failed rows")
    rows = matrix.get("rows")
    if not isinstance(rows, list) or len(rows) != 6:
        raise DeliveryReportError("supervised matrix does not contain six rows")
    for row in rows:
        if not isinstance(row, Mapping) or row.get("status") != "complete" or not isinstance(row.get("fixed_test"), Mapping):
            raise DeliveryReportError("a supervised matrix row lacks complete fixed-test evidence")
    return rows


def _validate_student(student: Mapping[str, Any]) -> Mapping[str, Any]:
    if student.get("artifact_type") != "student_fixed_test_evaluation" or student.get("exploratory") is not True:
        raise DeliveryReportError("Student evidence is not the expected exploratory fixed-test artifact")
    if not isinstance(student.get("metrics"), Mapping) or not isinstance(student.get("checkpoint"), Mapping):
        raise DeliveryReportError("Student metrics/checkpoint evidence is missing")
    if not student.get("run_id") or not student.get("split_fingerprint"):
        raise DeliveryReportError("Student run or split identity is missing")
    return student


def _validate_open_world(open_world: Mapping[str, Any]) -> Mapping[str, Any]:
    if open_world.get("artifact_type") != "post_student_open_world_discovery":
        raise DeliveryReportError("open-world artifact type is unsupported")
    if not isinstance(open_world.get("metrics"), Mapping) or not isinstance(open_world.get("split"), Mapping):
        raise DeliveryReportError("open-world metrics/split evidence is missing")
    return open_world


def _build_documents(summary: Mapping[str, Any], destination: Path) -> tuple[Path, Path]:
    from docx import Document
    from docx.shared import Pt

    matrix = summary["formal_matrix"]
    student = summary["student"]
    open_world = summary["open_world"]
    docx_path = destination / "delivery_evidence_report.docx"
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.add_heading("Semi-Supervised Fruit Object Detection", 0)
    doc.add_paragraph("Evidence-bound research and demonstration delivery report")
    doc.add_heading("Scope", level=1)
    doc.add_paragraph(
        "Known runtime classes are Apple, Banana, Orange, Strawberry and Pineapple. "
        "The PySide6 demonstrator is offline file/image based; camera control is not part of this delivery. "
        "The Student and open-world results are exploratory evidence, not a formal 0.80 acceptance claim."
    )
    doc.add_heading("Exploratory Student result", level=1)
    sm = student["metrics"]
    doc.add_paragraph(
        f"Run {student['run_id']} fixed-test mAP@0.5 {_metric_text(sm.get('map50'))}, "
        f"mAP@0.5:0.95 {_metric_text(sm.get('map50_95'))}, Precision {_metric_text(sm.get('precision'))}, "
        f"Recall {_metric_text(sm.get('recall'))}, F1 {_metric_text(sm.get('f1'))}."
    )
    doc.add_paragraph(f"Protected test split fingerprint: {student['split_fingerprint']}.")
    doc.add_heading("Formal supervised reference matrix", level=1)
    doc.add_paragraph(
        "Six canonical configurations completed with one fixed split protocol. "
        "The 100%-label row is a supervised upper-bound diagnostic; it does not constitute a formal SSOD acceptance result."
    )
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, ("Label budget", "Seed", "mAP50", "mAP50-95", "Precision", "Recall")):
        cell.text = title
    for row in matrix["rows"]:
        metrics = row["fixed_test"]
        cells = table.add_row().cells
        values = (f"{row['label_budget_percent']}%", str(row["seed"]), metrics.get("map50"), metrics.get("map50_95"), metrics.get("precision"), metrics.get("recall"))
        for cell, value in zip(cells, values):
            cell.text = _metric_text(value) if isinstance(value, (int, float)) else str(value)
    gate = matrix["upper_bound_gate"]
    doc.add_paragraph(f"Matrix upper-bound diagnostic: mAP@0.5 {_metric_text(gate.get('fixed_test_map50'))}; status {gate.get('status')}. No acceptance claim is made.")
    doc.add_heading("Five-class-outside discovery", level=1)
    discovery, holdout = open_world["metrics"]["discovery"], open_world["metrics"]["holdout"]
    doc.add_paragraph(
        f"The independent pool contains {open_world['split']['image_count']} images and six novel categories: "
        f"{', '.join(open_world['novel_categories_for_protected_evaluation'])}. "
        f"Discovery purity/NMI/ARI {_metric_text(discovery.get('purity'))}/{_metric_text(discovery.get('nmi'))}/{_metric_text(discovery.get('ari'))}; "
        f"holdout {_metric_text(holdout.get('purity'))}/{_metric_text(holdout.get('nmi'))}/{_metric_text(holdout.get('ari'))}."
    )
    novelty = open_world["novelty"]
    doc.add_paragraph(
        f"Unknown candidates: {novelty['candidate_count']} at threshold {_metric_text(novelty['threshold'])}; "
        "cluster names are post-hoc evaluation mappings and no new runtime class ID was added."
    )
    doc.add_heading("Evidence and limitations", level=1)
    doc.add_paragraph(
        "All numbers above are loaded from the sealed JSON files listed in the package manifest. "
        "The historical 0.85 upper-bound diagnostic is retained for transparency; it is not used to stop the customer-authorized exploratory path. "
        "Known-test false-positive rate was not measured because no known-test list was supplied."
    )
    doc.save(docx_path)

    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages

    pdf_path = destination / "delivery_evidence_report.pdf"
    with PdfPages(pdf_path) as pdf:
        fig = plt.figure(figsize=(8.27, 11.69))
        lines = [
            "Semi-Supervised Fruit Object Detection",
            "Evidence-bound research and demonstration delivery report",
            "",
            "Known classes: Apple, Banana, Orange, Strawberry, Pineapple",
            f"Student fixed-test mAP50: {_metric_text(sm.get('map50'))}",
            f"Formal matrix: {matrix['summary']['complete_runs']}/6 complete",
            f"100% supervised mAP50: {_metric_text(gate.get('fixed_test_map50'))}",
            "No formal 0.80 acceptance claim",
            "",
            "Open-world image-level discovery",
            f"Novel pool: {open_world['split']['image_count']} images; six categories",
            f"Discovery purity/NMI/ARI: {_metric_text(discovery.get('purity'))}/{_metric_text(discovery.get('nmi'))}/{_metric_text(discovery.get('ari'))}",
            f"Holdout purity/NMI/ARI: {_metric_text(holdout.get('purity'))}/{_metric_text(holdout.get('nmi'))}/{_metric_text(holdout.get('ari'))}",
            f"Unknown candidates: {novelty['candidate_count']} at threshold {_metric_text(novelty['threshold'])}",
            "Runtime registry remains five classes; camera is not included.",
        ]
        fig.text(0.08, 0.94, lines[0], fontsize=18, weight="bold")
        fig.text(0.08, 0.90, "\n".join(lines[1:]), fontsize=10, va="top", family="DejaVu Sans")
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)
    return docx_path, pdf_path


def build_delivery_report(*, matrix_path: Path, student_path: Path, open_world_path: Path, gui_metadata_path: Path, output: Path) -> Path:
    destination = output.resolve(strict=False)
    if destination.exists():
        raise DeliveryReportError(f"output already exists; refusing to overwrite: {destination}")
    matrix = _load(matrix_path.resolve(strict=True), "supervised matrix")
    rows = _validate_matrix(matrix)
    student = _validate_student(_load(student_path.resolve(strict=True), "Student fixed-test evidence"))
    open_world = _validate_open_world(_load(open_world_path.resolve(strict=True), "open-world evidence"))
    gui_metadata = _load(gui_metadata_path.resolve(strict=True), "GUI metadata")
    if gui_metadata.get("camera_enabled") is not False or gui_metadata.get("open_world_enabled") is not False:
        raise DeliveryReportError("GUI metadata does not prove camera/open-world runtime modes are disabled")
    summary = {
        "schema_version": "1.0",
        "protocol": "fruit_ssod_delivery_report_v1",
        "acceptance_claim": "none",
        "known_classes": ["Apple", "Banana", "Orange", "Strawberry", "Pineapple"],
        "formal_matrix": {"protocol": matrix["protocol"], "summary": matrix["summary"], "rows": rows, "upper_bound_gate": matrix["upper_bound_gate"]},
        "student": dict(student),
        "open_world": dict(open_world),
        "gui_metadata": dict(gui_metadata),
        "evidence": {
            "matrix": _evidence(matrix_path, "supervised matrix"),
            "student": _evidence(student_path, "Student fixed-test evidence"),
            "open_world": _evidence(open_world_path, "open-world evidence"),
            "gui_metadata": _evidence(gui_metadata_path, "GUI metadata"),
        },
    }
    destination.parent.mkdir(parents=True, exist_ok=True)
    temp = Path(tempfile.mkdtemp(prefix=f".{destination.name}.tmp-", dir=destination.parent))
    try:
        (temp / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8")
        lines = [
            "# Evidence-bound delivery report",
            "",
            "This package combines the completed supervised reference matrix with the customer-authorized exploratory Student and image-level open-world evidence. It makes no formal 0.80 acceptance claim.",
            "",
            f"Student fixed-test mAP@0.5: `{float(student['metrics']['map50']):.10f}`.",
            f"Formal matrix rows: `{matrix['summary']['complete_runs']}` complete, `{matrix['summary']['failed_runs']}` failed.",
            f"100% supervised fixed-test mAP@0.5: `{float(matrix['upper_bound_gate']['fixed_test_map50']):.10f}`.",
            f"Open-world pool: `{open_world['split']['image_count']}` images, `{len(open_world['novel_categories_for_protected_evaluation'])}` novel categories.",
            "",
            "All source JSON paths and SHA-256 digests are recorded in summary.json and manifest.json.",
        ]
        (temp / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
        temp.replace(destination)
        docx, pdf = _build_documents(summary, destination)
        manifest = {
            "protocol": summary["protocol"],
            "summary": _evidence(destination / "summary.json", "summary"),
            "readme": _evidence(destination / "README.md", "README"),
            "report_docx": _evidence(docx, "DOCX report"),
            "report_pdf": _evidence(pdf, "PDF report"),
        }
        (destination / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    except Exception:
        shutil.rmtree(temp, ignore_errors=True)
        raise
    return destination


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--matrix", type=Path, required=True)
    parser.add_argument("--student-test", type=Path, required=True)
    parser.add_argument("--open-world", type=Path, required=True)
    parser.add_argument("--gui-metadata", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    result = build_delivery_report(matrix_path=args.matrix, student_path=args.student_test, open_world_path=args.open_world, gui_metadata_path=args.gui_metadata, output=args.output)
    print(json.dumps({"output": str(result)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
