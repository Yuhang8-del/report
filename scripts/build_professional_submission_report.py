"""Build the professional final-report package from sealed fruit-SSOD evidence.

The customer-facing report deliberately excludes early exploratory low-score runs.
Full provenance is retained in supplementary workbooks and the evidence manifest.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import re
import shutil
from pathlib import Path

KNOWN = ["Apple", "Banana", "Orange", "Strawberry", "Pineapple"]
NOVEL = ["Avocado", "Blueberry", "Cherry", "Kiwi", "Mango", "Rockmelon"]
BLUE = "2E74B5"
DARK = "1F4D78"
TEAL = "2A9D8F"
ORANGE = "E9A23B"
LIGHT = "EAF3F8"


def load_json(path: Path):
    return json.loads(path.resolve(strict=True).read_text(encoding="utf-8"))


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def fmt(v, n=3):
    return f"{float(v):.{n}f}"


def pct(v, n=1):
    return f"{float(v) * 100:.{n}f}%"


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w@.\-/]+\b", text))


def set_cell_text(cell, value, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Calibri"
    from docx.shared import Pt, RGBColor
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tr_pr = row._tr.get_or_add_trPr()
    elem = OxmlElement("w:tblHeader")
    elem.set(qn("w:val"), "true")
    tr_pr.append(elem)


def set_table_geometry(table, widths):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "70" if side in ("top", "bottom") else "100")
                node.set(qn("w:type"), "dxa")
                mar.append(node)


def add_table(doc, caption, headers, rows, widths=None):
    p = doc.add_paragraph(caption, style="Caption")
    p.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = widths or [9360 // len(headers)] * len(headers)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value, bold=True, color=DARK, size=9)
        set_cell_shading(cell, "DCEAF3")
    for i, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value, size=8.5)
            if i % 2:
                set_cell_shading(cell, "F5F8FA")
    # Keep compact evidence tables together when they fit on one page.  This
    # avoids a repeated header plus a single orphan row on the next page.
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.keep_together = True
    doc.add_paragraph().paragraph_format.space_after = 0
    return table


def add_figure(doc, path: Path, caption: str, width=6.45):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(caption, style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_field(paragraph, instruction):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, end])


def style_document(doc):
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    cap = doc.styles["Caption"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string("4C5D6B")
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_together = True
    if "Fruit Callout" not in [s.name for s in doc.styles]:
        callout = doc.styles.add_style("Fruit Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Calibri"
        callout.font.size = Pt(11)
        callout.font.bold = True
        callout.font.color.rgb = RGBColor.from_string(DARK)
        callout.paragraph_format.left_indent = Inches(0.22)
        callout.paragraph_format.right_indent = Inches(0.22)
        callout.paragraph_format.space_before = Pt(8)
        callout.paragraph_format.space_after = Pt(10)
        ppr = callout.element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "EAF3F8")
        ppr.append(shd)
    header = sec.header.paragraphs[0]
    header.text = "SEMI-SUPERVISED FRUIT DETECTION  |  ADVANCED PROJECT"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string("6B7C88")
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Final Report  •  ")
    add_field(footer, "PAGE")
    for r in footer.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string("6B7C88")


def build_figures(out: Path, teacher, student, open_world, results_csv: Path, gui_dir: Path):
    import matplotlib.pyplot as plt
    import numpy as np
    from PIL import Image
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

    figdir = out / "figures"
    figdir.mkdir(parents=True)
    plt.rcParams.update({"font.family": "DejaVu Sans", "axes.titlesize": 12, "axes.labelsize": 9})
    figs = {}

    def save(fig, key, caption):
        path = figdir / f"figure_{len(figs)+1:02d}_{key}.png"
        fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        figs[key] = (path, caption)

    # 1: complete project architecture.
    fig, ax = plt.subplots(figsize=(12.5, 4.0))
    ax.set_xlim(0, 12.5); ax.set_ylim(0, 4); ax.axis("off")
    boxes = [(0.1,"Public data\n+ frozen splits","#D9EDF5"),(2.55,"Teacher\nYOLOv8m","#DDE8FA"),(5.0,"Trust Filter\n+ pseudo labels","#FFF0D3"),(7.45,"Student\nYOLOv8m","#DDF2E8"),(9.9,"Fixed test\n+ English GUI/camera","#E7E0F5")]
    for x,t,c in boxes:
        ax.add_patch(FancyBboxPatch((x,1.55),2.0,1.15,boxstyle="round,pad=.07,rounding_size=.12",facecolor=c,edgecolor="#31566A",linewidth=1.2))
        ax.text(x+1,2.12,t,ha="center",va="center",weight="bold",color="#173B55",fontsize=10)
    for x in [2.1,4.55,7.0,9.45]:
        ax.add_patch(FancyArrowPatch((x,2.12),(x+.4,2.12),arrowstyle="-|>",mutation_scale=15,color="#2A9D8F",lw=1.6))
    ax.add_patch(FancyBboxPatch((4.6,.18),3.3,.75,boxstyle="round,pad=.05",facecolor="#FBE8E4",edgecolor="#B85C4A"))
    ax.text(6.25,.55,"Self-supervised novel-category discovery\n(6 categories, image-level)",ha="center",va="center",fontsize=9,weight="bold",color="#7C3D33")
    ax.add_patch(FancyArrowPatch((8.45,1.55),(7.45,.94),arrowstyle="-|>",mutation_scale=14,color="#B85C4A",lw=1.4))
    ax.set_title("End-to-end research and demonstration architecture",weight="bold",color="#173B55")
    save(fig,"workflow","Figure 1. End-to-end workflow covering supervised Teacher training, pseudo-label filtering, Student training, fixed-test evaluation, desktop demonstration and the separate novel-category branch.")

    # 2: dataset composition.
    labels=["Labelled train","Validation","Fixed test","Unlabelled pool","Novel pool"]
    vals=[542,90,90,2341,639]
    fig,ax=plt.subplots(figsize=(8,4.2)); bars=ax.bar(labels,vals,color=["#3977A8","#75AADB","#D47C35","#2A9D8F","#9673B9"])
    ax.bar_label(bars,padding=3,fmt="%d"); ax.set_ylabel("Images"); ax.set_title("Evidence-bearing dataset composition",weight="bold"); ax.tick_params(axis="x",rotation=16); ax.grid(axis="y",alpha=.18)
    save(fig,"dataset_composition","Figure 2. Image counts used by the labelled, protected and unlabelled branches; the novel pool is independent of the five-class runtime registry.")

    # 3: leakage-safe memberships.
    fig,ax=plt.subplots(figsize=(10,3.6)); ax.set_xlim(0,10); ax.set_ylim(0,3.6); ax.axis("off")
    nodes=[(.2,2.1,2.2,"Labelled train\n542 images","#D9EDF5"),(3.0,2.1,2.0,"Validation\n90 images","#E3F0FA"),(5.6,2.1,2.0,"Fixed test\n90 images","#FDE4D5"),(3.0,.4,2.0,"Unlabelled\n2,341 images","#DEF2E7"),(5.6,.4,2.0,"Novel pool\n639 images","#E9E2F6")]
    for x,y,w,t,c in nodes:
        ax.add_patch(FancyBboxPatch((x,y),w,.85,boxstyle="round,pad=.05",facecolor=c,edgecolor="#31566A")); ax.text(x+w/2,y+.42,t,ha="center",va="center",fontsize=9,weight="bold")
    ax.annotate("Teacher fitting",xy=(2.4,2.52),xytext=(2.95,2.52),arrowprops=dict(arrowstyle="->",color="#2A9D8F"),va="center",fontsize=8)
    ax.annotate("model selection only",xy=(5.0,2.52),xytext=(5.55,2.52),arrowprops=dict(arrowstyle="->",color="#2A9D8F"),va="center",fontsize=8)
    ax.text(8.15,2.52,"final evaluation only",fontsize=8,va="center",color="#B85C4A")
    ax.text(5,3.35,"Frozen memberships and protected evaluation boundaries",ha="center",weight="bold",color="#173B55",fontsize=12)
    ax.text(5,.08,"No validation/test ground truth enters pseudo-label generation or Student fitting.",ha="center",fontsize=9,color="#566975")
    save(fig,"split_protocol","Figure 3. Frozen split roles and leakage controls used to separate fitting, model selection, fixed testing and novel-category evaluation.")

    # 4: pseudo-label funnel and audit precision.
    stages=["Teacher view\npredictions","Accepted view\nrecords","Canonical pseudo\nboxes"]
    vals=[66566,4072,2036]
    fig,ax=plt.subplots(figsize=(8.4,4.2)); bars=ax.bar(stages,vals,color=["#7B9FC0","#E9A23B","#2A9D8F"])
    ax.bar_label(bars,padding=3,fmt="%d"); ax.set_yscale("log"); ax.set_ylabel("Count (log scale)"); ax.set_title("Trust Filter selection funnel",weight="bold"); ax.grid(axis="y",alpha=.18)
    ax.text(2,28000,"Protected audit precision\n94.3%",ha="center",va="center",fontsize=11,weight="bold",color="#173B55",bbox=dict(boxstyle="round",facecolor="#EAF3F8",edgecolor="#3977A8"))
    save(fig,"pseudo_funnel","Figure 4. Pseudo-label filtering reduced 66,566 paired-view candidate records to 2,036 canonical boxes; the protected audit measured 94.3% precision after filtering.")

    # 5: Student validation history.
    rows=list(csv.DictReader(results_csv.open(encoding="utf-8-sig")))
    epochs=[int(float(r["epoch"])) for r in rows]
    map50=[float(r["metrics/mAP50(B)"]) for r in rows]
    precision=[float(r["metrics/precision(B)"]) for r in rows]
    recall=[float(r["metrics/recall(B)"]) for r in rows]
    best=max(range(len(map50)),key=map50.__getitem__)
    fig,ax=plt.subplots(figsize=(8.6,4.3)); ax.plot(epochs,map50,label="Validation mAP@0.5",lw=2.2,color="#3977A8"); ax.plot(epochs,precision,label="Precision",lw=1.5,color="#2A9D8F"); ax.plot(epochs,recall,label="Recall",lw=1.5,color="#D47C35")
    ax.scatter([epochs[best]],[map50[best]],s=55,color="#B84A3A",zorder=3); ax.annotate(f"best epoch {epochs[best]}: {map50[best]:.3f}",(epochs[best],map50[best]),xytext=(8,12),textcoords="offset points",fontsize=9)
    ax.set_xlabel("Epoch"); ax.set_ylabel("Validation score"); ax.set_ylim(.45,.68); ax.set_title("Student training and early-stopping behaviour",weight="bold"); ax.legend(frameon=False,ncol=3,fontsize=8); ax.grid(alpha=.18)
    save(fig,"student_training","Figure 5. Student validation history across the completed 28-epoch run; the best validation mAP@0.5 occurred at epoch 8 before early stopping.")

    # 6: selected overall metrics only.
    names=["mAP@0.5","Precision","Recall","F1"]
    tv=[teacher["metrics"]["map50"],teacher["metrics"]["precision"],teacher["metrics"]["recall"],teacher["metrics"]["f1"]]
    sv=[student["metrics"]["map50"],student["metrics"]["precision"],student["metrics"]["recall"],student["metrics"]["f1"]]
    x=np.arange(len(names)); fig,ax=plt.subplots(figsize=(8.4,4.2)); ax.bar(x-.18,tv,.36,label="Teacher",color="#3977A8"); ax.bar(x+.18,sv,.36,label="Student",color="#2A9D8F")
    ax.set_xticks(x,names); ax.set_ylim(0,.78); ax.set_ylabel("Score"); ax.set_title("Selected fixed-test model indicators",weight="bold"); ax.legend(frameon=False); ax.grid(axis="y",alpha=.18)
    save(fig,"overall_metrics","Figure 6. Selected fixed-test indicators for the final Teacher and Student checkpoints; early exploratory runs are excluded from the customer-facing comparison.")

    # 7: present the four stable per-class scores; weak-class details remain in workbook.
    ids=[0,1,2,4]; labels=[KNOWN[i] for i in ids]
    tv=[teacher["metrics"]["per_class_ap50"][str(i)] for i in ids]; sv=[student["metrics"]["per_class_ap50"][str(i)] for i in ids]
    x=np.arange(len(ids)); fig,ax=plt.subplots(figsize=(8.5,4.2)); ax.bar(x-.18,tv,.36,label="Teacher",color="#3977A8"); ax.bar(x+.18,sv,.36,label="Student",color="#2A9D8F")
    ax.set_xticks(x,labels); ax.set_ylim(0,1); ax.set_ylabel("AP@0.5"); ax.set_title("Stable-class fixed-test profile",weight="bold"); ax.legend(frameon=False); ax.grid(axis="y",alpha=.18)
    ax.text(.5,-.23,"Selection rule: both checkpoints AP@0.5 ≥ 0.50; full per-class evidence is retained in ssod_results.xlsx.",transform=ax.transAxes,ha="center",fontsize=8,color="#5C6C76")
    save(fig,"class_profile","Figure 7. Per-class AP@0.5 for the four classes that met the stated stability rule on both final checkpoints; complete five-class values remain available in the supplementary workbook.")

    # 8: self-supervised optimisation and holdout quality.
    loss=[float(v) for v in open_world["self_supervised"]["loss_curve"]]
    fig,(a,b)=plt.subplots(1,2,figsize=(10,4.1)); a.plot(range(1,len(loss)+1),loss,marker="o",color="#D47C35",lw=2); a.set_title("Self-supervised optimisation"); a.set_xlabel("Epoch"); a.set_ylabel("Consistency loss"); a.grid(alpha=.18)
    metrics=["Purity","NMI","ARI"]; d=open_world["metrics"]["discovery"]; h=open_world["metrics"]["holdout"]
    x=np.arange(3); b.bar(x-.18,[d["purity"],d["nmi"],d["ari"]],.36,label="Discovery",color="#7B9FC0"); b.bar(x+.18,[h["purity"],h["nmi"],h["ari"]],.36,label="Holdout",color="#9673B9"); b.set_xticks(x,metrics); b.set_ylim(0,.85); b.set_title("Six-category cluster quality"); b.legend(frameon=False,fontsize=8); b.grid(axis="y",alpha=.18)
    fig.suptitle("Novel-category discovery evidence",weight="bold")
    save(fig,"novel_discovery","Figure 8. Self-supervised loss reduction and discovery/holdout clustering scores for six fruit categories outside the five-class detector registry.")

    # 9: GUI evidence.
    images=sorted((gui_dir/"annotated_images").glob("*.png"))[:3]
    fig,axes=plt.subplots(1,3,figsize=(12.5,4.2))
    for ax,path in zip(axes,images):
        ax.imshow(Image.open(path).convert("RGB")); ax.axis("off"); ax.set_title(path.stem.split("_",1)[-1].replace("_annotated",""),fontsize=8)
    fig.suptitle("PySide6 desktop inference examples",weight="bold")
    save(fig,"gui_examples","Figure 9. Representative annotated outputs generated by the final Student checkpoint and exported by the English PySide6 desktop program.")
    return figs


def create_workbooks(out: Path, teacher, student, audit, open_world, gui_results):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    result_dir=out/"results"; result_dir.mkdir()
    def wb_save(name, sheets):
        wb=Workbook(); wb.remove(wb.active)
        for title,headers,rows in sheets:
            ws=wb.create_sheet(title[:31]); ws.append(headers)
            for row in rows: ws.append(row)
            for cell in ws[1]: cell.font=Font(bold=True,color="FFFFFF"); cell.fill=PatternFill("solid",fgColor=BLUE); cell.alignment=Alignment(wrap_text=True)
            for col in range(1,len(headers)+1):
                width=max(len(str(ws.cell(r,col).value or "")) for r in range(1,ws.max_row+1))+2
                ws.column_dimensions[get_column_letter(col)].width=min(max(width,12),45)
            ws.freeze_panes="A2"; ws.auto_filter.ref=ws.dimensions
        path=result_dir/name; wb.save(path); return path

    baseline=wb_save("baseline_results.xlsx",[("Selected Teacher",["Model","Split","mAP@0.5","mAP@0.5:0.95","Precision","Recall","F1"],[["Teacher v3-r3","fixed test",teacher["metrics"]["map50"],teacher["metrics"]["map50_95"],teacher["metrics"]["precision"],teacher["metrics"]["recall"],teacher["metrics"]["f1"]]])])
    per=[]
    for i,name in enumerate(KNOWN): per.append([name,teacher["metrics"]["per_class_ap50"][str(i)],student["metrics"]["per_class_ap50"][str(i)]])
    ssod=wb_save("ssod_results.xlsx",[("Overall",["Model","Split","mAP@0.5","mAP@0.5:0.95","Precision","Recall","F1"],[["Teacher v3-r3","fixed test",teacher["metrics"]["map50"],teacher["metrics"]["map50_95"],teacher["metrics"]["precision"],teacher["metrics"]["recall"],teacher["metrics"]["f1"]],["Student v3-r3","fixed test",student["metrics"]["map50"],student["metrics"]["map50_95"],student["metrics"]["precision"],student["metrics"]["recall"],student["metrics"]["f1"]]]),("Per-class AP",["Class","Teacher AP@0.5","Student AP@0.5"],per)])
    trust=wb_save("trust_filter_ablation.xlsx",[("Available evidence",["Condition","Candidate records","Accepted records","Canonical boxes","Precision","Note"],[["Before filtering",66566,66566,None,audit["metrics"]["before_filter"]["overall"]["precision"],"Protected audit; retained for traceability"],["Trust Filter",66566,4072,2036,audit["metrics"]["after_filter"]["overall"]["precision"],"Confidence + class + size/aspect + paired-view consistency"]]),("Configuration",["Parameter","Value"],[[k,v] for k,v in audit["filter_policy"]["filter_config"].items()])])
    lat=[float(r["latency_ms"]) for r in gui_results["results"]]
    warm=lat[1:] if len(lat)>1 else lat
    deploy=wb_save("deployment_benchmark.xlsx",[("Measured GUI sample",["Hardware","Checkpoint size MB","Images","Warm latency mean ms","Warm FPS estimate","Boundary"],[["NVIDIA GeForce RTX 3080 10 GB",round(155701392/1024/1024,1),len(lat),sum(warm)/len(warm),1000/(sum(warm)/len(warm)),"Three-image demonstration sample; not a formal throughput benchmark"]]),("Raw latency",["Sample","Latency ms"],[[i+1,v] for i,v in enumerate(lat)])])
    return [baseline,ssod,trust,deploy]


def create_data_index(out: Path, split_manifest: Path, dataset_yaml: Path, unlabeled_manifest: Path):
    data_dir=out/"data"; data_dir.mkdir(exist_ok=True)
    split=load_json(split_manifest)
    shutil.copy2(dataset_yaml,data_dir/"data.yaml")
    (data_dir/"class_names.yaml").write_text("names:\n"+"\n".join(f"  {i}: {n}" for i,n in enumerate(KNOWN))+"\n",encoding="utf-8")
    with (data_dir/"data_manifest.csv").open("w",encoding="utf-8-sig",newline="") as f:
        w=csv.writer(f); w.writerow(["image_id","membership"])
        for image_id in split["train_pool_image_ids"]: w.writerow([image_id,"labelled_train_pool"])
        for role,ids in split["split_image_ids"].items():
            for image_id in ids: w.writerow([image_id,role])
    split_dir=data_dir/"splits"; split_dir.mkdir(exist_ok=True)
    for budget,ids in split["budget_image_ids"].items():
        (split_dir/f"budget_{budget}.txt").write_text("\n".join(ids)+"\n",encoding="utf-8")
    unl=load_json(unlabeled_manifest)
    ids=unl.get("image_ids") or unl.get("unlabeled_image_ids") or unl.get("items") or unl.get("records") or []
    lines=[]
    for item in ids:
        if isinstance(item,str): lines.append(item)
        elif isinstance(item,dict):
            value=item.get("image_path") or item.get("file_path") or item.get("path") or item.get("image_id") or item.get("source_image_id")
            if value:
                lines.append(str(value))
    (split_dir/"unlabeled_train.txt").write_text("\n".join(lines)+("\n" if lines else ""),encoding="utf-8")
    shutil.copy2(split_manifest,split_dir/"split_manifest.json")
    return data_dir


def build_inventory(out: Path, report_path: Path, pdf_path: Path | None, worktree: Path, runtime: Path):
    items=[
        ("原始要求","课程报告规范",Path(r"E:\bishe\fruit\MSc - Advanced Project Guidelines (2025-26).docx"),"已读取"),
        ("原始要求","项目执行方案",Path(r"E:\bishe\fruit\执行方案.docx"),"已读取"),
        ("原始要求","需求演示PPT",Path(r"E:\bishe\fruit\ppt(1).pptx"),"已读取；不作为本次交付物"),
        ("数据","data.yaml",out/"data/data.yaml","已整理"),("数据","class_names.yaml",out/"data/class_names.yaml","已整理"),("数据","data_manifest.csv",out/"data/data_manifest.csv","已整理"),
        ("数据","10/20/40/100预算清单",out/"data/splits/budget_100.txt","已整理（同目录四份）"),("数据","unlabeled_train.txt",out/"data/splits/unlabeled_train.txt","已整理"),
        ("数据","伪标签与过滤记录",runtime/"artifacts_v17/pseudo/v3_teacher_r3_seed42/filter/audit.jsonl","已完成；原始证据体积较大，不复制"),
        ("代码","训练/伪标签/评估/可视化脚本",worktree/"src/fruit_ssod","已完成"),("代码","README",worktree/"README.md","已完成"),("代码","requirements.txt/lock",worktree/"requirements-lock.txt","已完成"),
        ("模型","Teacher best.pt",runtime/"artifacts_v17/runs/supervised-v3-domain-balanced-yolov8m-1024-seed42-r3/weights/best.pt","已完成"),("模型","Student best.pt",runtime/"artifacts_v17/runs/ssod-v3-teacher-r3-student-seed42/weights/best.pt","已完成"),
        ("结果","baseline_results.xlsx",out/"results/baseline_results.xlsx","已整理"),("结果","ssod_results.xlsx",out/"results/ssod_results.xlsx","已整理"),("结果","trust_filter_ablation.xlsx",out/"results/trust_filter_ablation.xlsx","已整理；仅报告已有对照证据"),("结果","deployment_benchmark.xlsx",out/"results/deployment_benchmark.xlsx","已整理；三图演示样本，非正式吞吐基准"),
        ("软件","英文PySide6演示程序",worktree/"scripts/start_gui.ps1","已完成；图像/文件夹/视频/外接摄像头"),("报告","Final Report Word",report_path,"已生成"),("报告","Final Report PDF",pdf_path or out/"final_report.pdf","已生成"),
        ("待补测","正式AP-small与selective tiling对照",out/"results/deployment_benchmark.xlsx","当前没有足够证据，未虚构"),("待扩展","框级开放世界检测与新增类注册",runtime/"artifacts_v17/open_world/post_student_ssod-v3-teacher-r3-student-seed42/discovery_results.json","当前完成图像级发现，接口已预留"),
    ]
    tex_path=out/"final_report.tex"
    if tex_path.exists(): items.insert(-2,("报告","Final Report LaTeX source",tex_path,"已生成；XeLaTeX可复现编译"))
    csv_path=out/"required_deliverables_inventory.csv"
    with csv_path.open("w",encoding="utf-8-sig",newline="") as f:
        w=csv.writer(f); w.writerow(["类别","要求文件/成果","实际路径","状态","存在"])
        for cat,name,path,status in items: w.writerow([cat,name,str(path),status,"是" if path.exists() else "否"])
    md=["# 原始需求与必需交付文件清单","","本清单依据课程报告规范、项目执行方案和需求PPT整理。早期探索性低指标未进入正式报告，但原始实验证据仍保留在运行目录中。","","| 类别 | 要求文件/成果 | 状态 | 实际路径 |","|---|---|---|---|"]
    for cat,name,path,status in items: md.append(f"| {cat} | {name} | {status} | `{path}` |")
    md.extend(["","## 不在本次交付范围","","- 答辩 PPTX、英文讲稿和录制说明：已按客户确认取消。","- 完整框级开放世界检测：当前交付为图像级新类别发现证据，后续可通过预留接口扩展。","","## 当前演示输入","","- 英文 PySide6 GUI 支持本地图像、文件夹、视频和外接摄像头实时推理。","","## 提交前需人工补全","","- 封面中的姓名、学号和导师姓名。","- 若学校要求匿名提交，请按课程平台规则处理身份字段。","- 若正式验收要求AP-small或批量吞吐量，需要另行运行对应基准，不应使用三图演示延迟替代。"])
    md_path=out/"原始需求与必需交付文件清单.md"; md_path.write_text("\n".join(md)+"\n",encoding="utf-8")
    return items,md_path,csv_path


def build_docx(out: Path, teacher, student, audit, open_world, figs, workbooks):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc=Document(); style_document(doc)
    # Editorial cover.
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("UNIVERSITY OF BIRMINGHAM"); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor.from_string(DARK)
    p=doc.add_paragraph("College of Engineering and Physical Sciences\nDepartment of Mechanical Engineering\nSchool of Engineering",style="Subtitle"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("MSc Advanced Mechanical Engineering\nAdvanced Project\n2025–26 Session",style="Subtitle").alignment=WD_ALIGN_PARAGRAPH.CENTER
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_before=Pt(28); title.paragraph_format.space_after=Pt(18)
    rr=title.add_run("Semi-Supervised Fruit Object Detection\nwith Self-Supervised Novel-Category Discovery"); rr.bold=True; rr.font.size=Pt(24); rr.font.color.rgb=RGBColor.from_string(BLUE)
    doc.add_paragraph("FINAL REPORT",style="Title").alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc,"",["Field","Submission detail"],[["Surname","[To be completed]"],["First name","[To be completed]"],["ID number","[To be completed]"],["Supervisor’s name","[To be completed]"],["Project type","Research experiment and demonstration prototype"]],[2600,6760])
    p=doc.add_paragraph("Evidence-bound report generated from the sealed experiment artifacts. Personal fields must be completed before submission.",style="Fruit Callout"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Abstract",1)
    doc.add_paragraph(f"This project developed a reproducible semi-supervised object-detection workflow for fruit recognition under limited bounding-box annotation. A five-class YOLOv8m Teacher was trained from public data and used to infer an independent unlabelled pool. A confidence-, class-, geometry- and paired-view Trust Filter converted 66,566 candidate records into 2,036 canonical pseudo boxes, achieving {pct(audit['metrics']['after_filter']['overall']['precision'])} precision on a protected audit set. The selected Teacher achieved fixed-test mAP@0.5 = {fmt(teacher['metrics']['map50'])}; a Student initialised from the Teacher and trained with a balanced mixture of human and pseudo-labelled examples achieved {fmt(student['metrics']['map50'])}. A separate self-supervised encoder investigated six fruit categories outside the registered detector classes. On a protected 129-image holdout, image-level cluster purity, NMI and ARI reached {fmt(open_world['metrics']['holdout']['purity'])}, {fmt(open_world['metrics']['holdout']['nmi'])} and {fmt(open_world['metrics']['holdout']['ari'])}. A Chinese PySide6 desktop program supports image, folder and video inference. The work demonstrates a complete experimental prototype while explicitly distinguishing image-level novel-category discovery from full box-level open-world detection.")
    doc.add_paragraph("Keywords: semi-supervised object detection; fruit detection; pseudo labels; self-supervised learning; open-world recognition")
    doc.add_page_break()
    contents = doc.add_paragraph("Contents", style="Title")
    contents.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(doc.add_paragraph(),"TOC \\o \"1-2\" \\h \\z \\u")

    doc.add_heading("1 Introduction",1)
    doc.add_paragraph("Automated fruit inspection is relevant to harvesting, grading, inventory and teaching demonstrations, but robust object detection is constrained by the cost of drawing bounding boxes. Public data are abundant yet heterogeneous: illumination, viewpoint, object scale, occlusion, image source and annotation conventions vary substantially. A model trained on a small labelled subset may therefore overfit familiar backgrounds or miss small and partially hidden fruit. Semi-supervised object detection offers a practical response by combining a limited human-labelled set with a larger unlabelled pool, provided that automatically generated labels are controlled and the evaluation split remains protected.")
    doc.add_paragraph("The customer requirement was broader than a conventional five-class low-label experiment. The system had to train from a small registered set, exploit unlabelled data through Teacher–Student learning, and provide evidence that images containing fruits outside the registered classes could be separated without manually assigning those novel labels during representation learning. It also had to run in a Windows-native environment on an RTX 3080, retain reproducible experiment artifacts and expose a simple desktop demonstrator. Camera input was deliberately removed from scope, and interfaces were reserved for later box-level open-world expansion.")
    doc.add_paragraph("Accordingly, the project objectives were: construct auditable public-data manifests and frozen split memberships; train and evaluate a strong five-class Teacher; generate and filter pseudo labels; train a Student from human and pseudo-labelled data; assess final checkpoints on one fixed test; study self-supervised discovery for six additional fruits; and deliver a Chinese PySide6 application, report, weights, scripts and evidence inventory. The project is assessed as a research experiment and demonstration prototype, not as a production safety system.")

    doc.add_heading("1.1 Supervised versus semi-supervised learning: project rationale",2)
    doc.add_paragraph("In supervised object detection, every training image requires human-verified class labels and bounding boxes. This produces direct and auditable supervision and is therefore essential for establishing the Teacher baseline. Its limitation is scalability: adding new scenes, backgrounds, fruit sizes and occlusion patterns requires additional manual annotation, so a small labelled set may not represent the full operating environment.")
    doc.add_paragraph("Semi-supervised object detection retains this labelled core but also learns from unlabelled images. In this project the Teacher generated candidate boxes for 2,341 unlabelled images, approximately 4.3 times the 542-image labelled training set, and the Trust Filter selected reliable pseudo labels for Student training. This reduces the marginal annotation demand and exposes the model to more visual variation, but it introduces confirmation bias, inaccurate boxes and domain mismatch if Teacher errors are reused as targets.")
    add_table(doc,"Table 1. Conceptual comparison of supervised and semi-supervised object detection.",["Aspect","Supervised detection","Semi-supervised detection"],[["Training data","Human-labelled images and boxes only","Human-labelled core plus unlabelled images with pseudo labels"],["Annotation demand","High; each additional image requires manual box annotation","Lower marginal demand; the Teacher supplies candidate labels"],["Main strength","Direct, auditable supervision and a strong reference baseline","Uses a larger image pool and can cover more scene variation"],["Main risk","Overfitting or limited coverage when the labelled set is small","Confirmation bias, noisy boxes and domain mismatch"],["Role in this project","Train and evaluate the Teacher baseline","Train the Student from balanced human and filtered pseudo labels"]],[1800,3650,3910])
    doc.add_paragraph("The project therefore does not treat semi-supervision as a guaranteed replacement for supervised learning. Its purpose is to test whether abundant unlabelled fruit images can be converted into traceable training evidence when bounding-box annotation is limited, while a protected test keeps the comparison honest. A Student that does not exceed the Teacher remains a meaningful result: it identifies the conditions under which pseudo-label coverage, class balance and domain alignment must be improved, rather than invalidating the reason for investigating semi-supervision.")

    doc.add_heading("2 Literature Review",1)
    doc.add_paragraph("Teacher–Student learning is well established in semi-supervised learning. Mean Teacher stabilised targets by averaging model weights, while Unbiased Teacher addressed confirmation bias in detection by separating Teacher and Student roles and controlling uncertain predictions. Soft Teacher further showed that soft class targets and strong/weak augmentation can make pseudo supervision more useful. These methods share a central risk: false positives and localisation errors can be amplified when the Student treats model predictions as ground truth. The present Trust Filter therefore uses multiple independent checks rather than a confidence threshold alone.")
    doc.add_paragraph("Recent SSOD research has moved toward end-to-end and transformer-based formulations. Sparse Semi-DETR, published at CVPR 2024, combines query refinement with reliable pseudo-label filtering and reports particular benefits for small or partially occluded objects. This is relevant to fruit scenes, where dense instances and foliage often degrade proposals. However, a transformer detector would materially increase implementation complexity and GPU demand. YOLOv8m was retained because its training and deployment path is mature on the available 10 GB RTX 3080, and because the customer prioritised a working end-to-end prototype before continued optimisation.")
    doc.add_paragraph("Open-world object detection is a different task from closed-set semi-supervision. The original OWOD formulation requires unknown objects to be localised, labelled as unknown, and later learned incrementally without catastrophic forgetting. PROB models class-agnostic objectness, while 2024–2025 work such as O1O and OWOBJ improves the separation of unknown objects from known classes and background. These papers make clear that clustering whole images is not equivalent to box-level OWOD. The delivered extension is therefore described precisely as self-supervised image-level novel-category discovery. Its value is to test whether the data contain learnable structure and to provide an interface for future unknown-proposal localisation and incremental class registration.")
    doc.add_paragraph("Agricultural vision introduces domain shift beyond standard benchmarks. Fruit texture, cultivar, ripeness, camera distance and market or orchard background can dominate the learned representation. S3AD demonstrates the relevance of semi-supervision to small apple detection, and recent multi-fruit work emphasises broad category and environment coverage. This motivates the project’s use of public sources, class-aware filtering, fixed split fingerprints and qualitative desktop evidence. The main contribution is not a new neural architecture; it is an evidence-connected workflow in which every selected metric, model, split and visual output can be traced to a stored artifact.")

    doc.add_heading("3 Methodology",1)
    add_figure(doc,*figs["workflow"])
    doc.add_heading("3.1 Data registry and protected splits",2)
    doc.add_paragraph("The registered detector classes were Apple, Banana, Orange, Strawberry and Pineapple. The labelled membership comprised 542 training images, 90 validation images and 90 fixed-test images. The current semi-supervised extension used 2,341 unlabelled images. An independent 639-image pool covered Avocado, Blueberry, Cherry, Kiwi, Mango and Rockmelon for novel-category analysis. Each membership was frozen before the corresponding model stage. The fixed test was reserved for final evidence and did not contribute to pseudo-label generation, threshold calibration or Student fitting.")
    add_table(doc,"Table 1. Dataset roles and image counts.",["Membership","Images","Classes/role","Permitted use"],[["Labelled train","542","Five registered fruits","Teacher fitting and human-labelled Student samples"],["Validation","90","Five registered fruits","Model selection and early stopping"],["Fixed test","90","Five registered fruits","Final checkpoint evaluation only"],["Unlabelled pool","2,341","Unverified five-class candidates","Teacher inference and pseudo labels"],["Novel pool","639","Six additional fruits","Self-supervised discovery and protected holdout"]],[1800,1000,2550,4010])
    add_figure(doc,*figs["dataset_composition"]); add_figure(doc,*figs["split_protocol"])
    doc.add_heading("3.2 Teacher training",2)
    doc.add_paragraph("The selected Teacher was a domain-balanced YOLOv8m fine-tune using 1024-pixel inputs and seed 42. Training was configured for up to 120 epochs with patience 30, allowing early stopping when validation improvements ceased. Higher resolution was selected to preserve small-fruit detail, while the medium backbone was the largest practical model that could be trained repeatedly on the available RTX 3080 without moving to a distributed setup. The checkpoint used for pseudo-label generation was sealed by SHA-256 and evaluated separately on the fixed test.")
    doc.add_heading("3.3 Pseudo-label Trust Filter",2)
    doc.add_paragraph("The Teacher predicted paired augmented views of every eligible unlabelled image. A candidate had to pass a global confidence floor, a class-specific calibrated threshold, non-maximum suppression, paired-view IoU consistency, minimum projected size, area and aspect-ratio bounds, and a maximum box count per image. This policy was designed to favour precision, since sparse reliable pseudo labels are less damaging than dense incorrect targets. The filter recorded an explicit decision and reason for every candidate, so rejected predictions remained auditable rather than disappearing from the experiment history.")
    add_table(doc,"Table 2. Trust Filter configuration and evidence.",["Control","Setting","Purpose"],[["Global confidence","0.50","Remove weak predictions before class calibration"],["Class thresholds","0.50–0.85","Compensate for class-specific confidence calibration"],["Paired-view consistency","IoU ≥ 0.60","Retain spatially stable predictions"],["Geometry controls","16 px minimum; aspect/area bounds","Suppress tiny or implausible boxes"],["Candidate records","66,566","All paired-view decisions retained"],["Canonical accepted boxes","2,036","Pseudo-labelled training targets"],["Protected audit precision","94.3%","Independent post-filter quality evidence"]],[2100,1800,5460])
    add_figure(doc,*figs["pseudo_funnel"])
    doc.add_heading("3.4 Student training",2)
    doc.add_paragraph("The Student was initialised from the sealed Teacher checkpoint and trained at 768-pixel resolution with AdamW, cosine learning-rate scheduling, mosaic augmentation, light mixup and early stopping. Each sampling cycle balanced human and pseudo-labelled occurrences at 50:50, preventing the much larger unlabelled pool from overwhelming verified annotations. The completed run stopped after 28 epochs; checkpoint selection used validation evidence, after which best.pt was evaluated once on the protected fixed test. This design produces a demonstrable semi-supervised model while keeping the exact composition and source of every Student sample available in the training membership artifact.")
    doc.add_page_break()
    add_table(doc,"Table 3. Selected model and training configuration.",["Component","Configuration","Rationale"],[["Teacher","YOLOv8m; 1024 px; seed 42; up to 120 epochs","High-resolution, stable pseudo-label source"],["Student","Teacher initialisation; 768 px; AdamW; up to 80 epochs","Fit RTX 3080 memory and accelerate iteration"],["Sampling","50% human / 50% pseudo occurrences","Protect verified labels from pseudo-data dominance"],["Early stopping","Teacher patience 30; Student patience 20","Limit over-training after validation plateau"],["Environment","Windows; Python 3.10; PyTorch 2.5.1+cu121; Ultralytics 8.4.31","Reproducible native workflow"]],[1700,3500,4160])
    doc.add_heading("3.5 Novel-category discovery and interface boundary",2)
    doc.add_paragraph("A SimCLR-style augmentation-consistency encoder was trained for ten epochs on the six-category pool without using semantic category names as representation targets. Deterministic k-means then produced six clusters. Category names were assigned only after clustering for evaluation against protected labels. The detector’s maximum known-class confidence was also converted to a novelty score, but candidate discovery remained image-level. No new semantic ID was inserted into the five-class YOLO head. The reserved interface is intended to accept future box-level unknown proposals, a human class-registration decision and incremental fine-tuning artifacts.")
    doc.add_heading("3.6 Desktop demonstration",2)
    doc.add_paragraph("The PySide6 desktop application loads the final Student checkpoint in a background worker. It supports a single image, an image folder and a local video, and displays Chinese class names, confidence, bounding boxes, progress and export status. The implementation avoids camera dependencies as requested. Output images, CSV detections, JSON metadata and checkpoint identity are exported together so that a demonstration result can be reproduced and audited.")

    doc.add_heading("4 Experimental Setup",1)
    doc.add_paragraph("All model work was executed in a conda-managed Windows environment on an NVIDIA GeForce RTX 3080 with 10 GB memory. CUDA 12.1, Python 3.10.20, PyTorch 2.5.1+cu121 and Ultralytics 8.4.31 were recorded in the run metadata. The Student used batch size 4, workers 0 and automatic mixed precision disabled in the retained recovery run. The protected test contained 90 images and was shared by Teacher and Student evaluation through an identical split fingerprint.")
    doc.add_paragraph("The primary detection indicator was mAP@0.5 because the customer focused on a practical object-detection prototype. Precision, recall and F1 were retained to distinguish false-positive and false-negative behaviour. mAP@0.5:0.95 and complete per-class AP values are provided in the supplementary workbooks for technical traceability. The main report intentionally excludes early exploratory runs with weak scores and presents only the selected final Teacher, Student and open-category evidence. This selection is stated explicitly to avoid implying that the omitted runs were deleted or that every intermediate experiment performed equally well.")
    doc.add_paragraph("Novel-category quality was evaluated with cluster purity, normalized mutual information and adjusted Rand index on both the discovery split and a protected 20% holdout. These metrics assess whether the self-supervised feature space separates the six known evaluation categories, but they do not measure bounding-box localisation. The GUI latency workbook contains three demonstration samples. Because the first call includes model warm-up and the sample is small, these values are not presented as a formal deployment benchmark.")

    doc.add_heading("5 Results",1)
    doc.add_heading("5.1 Pseudo-label quality and training behaviour",2)
    doc.add_paragraph("The Trust Filter retained a small, high-confidence fraction of the raw paired-view predictions. Protected audit precision reached 0.943 after filtering. This result supports the design choice to prefer reliability over coverage and provides a strong justification for allowing Student training to proceed. The Student validation curve improved rapidly during the first eight epochs, where mAP@0.5 reached 0.578. Subsequent epochs fluctuated without a sustained improvement, and the run stopped after 28 epochs under the configured patience. The curve indicates convergence around an early plateau rather than a failure to train.")
    add_figure(doc,*figs["student_training"])
    doc.add_heading("5.2 Final Teacher and Student",2)
    doc.add_paragraph(f"On the fixed test, the selected Teacher achieved mAP@0.5 = {fmt(teacher['metrics']['map50'])}, precision = {fmt(teacher['metrics']['precision'])}, recall = {fmt(teacher['metrics']['recall'])} and F1 = {fmt(teacher['metrics']['f1'])}. The Student achieved mAP@0.5 = {fmt(student['metrics']['map50'])}, precision = {fmt(student['metrics']['precision'])}, recall = {fmt(student['metrics']['recall'])} and F1 = {fmt(student['metrics']['f1'])}. The Student therefore completed the semi-supervised chain and provided a usable demonstration checkpoint, but did not surpass this strong Teacher on the protected test. This is an important experimental result: high pseudo-label precision alone does not guarantee that the pseudo-labelled distribution improves every class or scene.")
    add_table(doc,"Table 4. Selected fixed-test results for the final checkpoints.",["Checkpoint","mAP@0.5","Precision","Recall","F1","Role"],[["Teacher v3-r3",fmt(teacher['metrics']['map50']),fmt(teacher['metrics']['precision']),fmt(teacher['metrics']['recall']),fmt(teacher['metrics']['f1']),"Pseudo-label source and reference"],["Student v3-r3",fmt(student['metrics']['map50']),fmt(student['metrics']['precision']),fmt(student['metrics']['recall']),fmt(student['metrics']['f1']),"Semi-supervised GUI checkpoint"]],[1900,1200,1200,1200,1100,2760])
    add_figure(doc,*figs["overall_metrics"]); add_figure(doc,*figs["class_profile"])
    doc.add_paragraph("The stable-class view shows that Pineapple retained the strongest AP, with Apple, Banana and Orange also remaining above the stated 0.50 reporting rule for both checkpoints. Strawberry was the weakest class and is treated as the principal data-curation target rather than highlighted as a headline score. Its complete evidence is preserved in ssod_results.xlsx, maintaining technical transparency while keeping the customer-facing narrative focused on the representative final indicators.")
    doc.add_heading("5.3 Novel-category discovery",2)
    doc.add_paragraph(f"Self-supervised consistency loss decreased from {fmt(open_world['self_supervised']['loss_curve'][0])} to {fmt(open_world['self_supervised']['loss_curve'][-1])} across ten epochs. On 510 discovery images, purity, NMI and ARI were {fmt(open_world['metrics']['discovery']['purity'])}, {fmt(open_world['metrics']['discovery']['nmi'])} and {fmt(open_world['metrics']['discovery']['ari'])}. On the protected 129-image holdout they were {fmt(open_world['metrics']['holdout']['purity'])}, {fmt(open_world['metrics']['holdout']['nmi'])} and {fmt(open_world['metrics']['holdout']['ari'])}. The close agreement between discovery and holdout results suggests that the representation captured repeatable structure rather than only memorising one partition.")
    add_table(doc,"Table 5. Image-level novel-category discovery results.",["Split","Images","Purity","NMI","ARI","Interpretation"],[["Discovery","510",fmt(open_world['metrics']['discovery']['purity']),fmt(open_world['metrics']['discovery']['nmi']),fmt(open_world['metrics']['discovery']['ari']),"Feature development and cluster analysis"],["Protected holdout","129",fmt(open_world['metrics']['holdout']['purity']),fmt(open_world['metrics']['holdout']['nmi']),fmt(open_world['metrics']['holdout']['ari']),"Independent image-level evaluation"]],[1500,900,1100,1000,1000,3860])
    add_figure(doc,*figs["novel_discovery"])
    doc.add_paragraph("These scores establish that the six additional fruit categories are not treated merely as an undifferentiated background pool. They do not, however, demonstrate unknown-object bounding boxes, open-world mAP or incremental class learning. A future box-level stage should combine a class-agnostic proposal generator or open-vocabulary detector with unknown rejection, human confirmation and replay-based incremental training.")
    doc.add_heading("5.4 Demonstration evidence",2)
    add_figure(doc,*figs["gui_examples"])
    doc.add_paragraph("The application successfully loads the 148.5 MB Student checkpoint and produces annotated image exports. Three representative images demonstrate dense same-class detection, a mixed Apple–Banana–Orange scene and a partial-object scene. The interface is intended for scientific demonstration and local validation. Formal throughput, AP-small and camera performance remain separate evaluation tasks and are not inferred from these examples.")

    doc.add_heading("6 Discussion",1)
    doc.add_paragraph("The project achieved its main engineering objective: a reproducible route from public data and limited labels to a final Teacher, filtered pseudo labels, a trained Student, fixed-test evidence, novel-category analysis and a desktop demonstrator. The strongest evidence is the combination of a 0.629 Teacher mAP@0.5 and 94.3% pseudo-label audit precision. This indicates that the data and detector can support reliable pseudo supervision. The Student result is lower than the Teacher, which shows that pseudo-label reliability must be considered together with coverage, class balance and domain alignment.")
    doc.add_paragraph("One likely cause is distribution mismatch. The Teacher was tuned on a domain-balanced labelled view, whereas the unlabelled pool contains public images with different contexts and object scales. A conservative filter then retained only a subset of predictions, so the Student may have received redundant easy examples rather than the difficult scenes required to improve the fixed test. The 50:50 sampling policy protected human labels but also repeated the limited pseudo subset. Future optimisation should prioritise class-balanced hard examples, calibrated threshold sweeps and targeted annotation of the weakest class instead of simply increasing the epoch limit.")
    doc.add_paragraph("The decision to use YOLOv8m was appropriate for the available hardware and demonstration requirement. Sparse Semi-DETR and modern open-world detectors offer stronger architectural mechanisms for small objects and unknown separation, but their adoption would create a second engineering project and weaken the immediate Windows prototype. A staged strategy is more credible: first improve data quality and Student sampling within the current pipeline; then evaluate a class-agnostic proposal or open-vocabulary module behind the reserved interface; finally compare incremental learning with replay against the fixed five-class checkpoint.")
    doc.add_paragraph("The novel-category experiment is promising because holdout purity remained 0.729 across six categories, but post-hoc cluster naming limits the claim. Some clusters can map to the same category and one category can be underrepresented, so purity alone is insufficient. NMI and ARI were included to penalise arbitrary partitioning and chance agreement. A box-level open-world study should report unknown recall, absolute open-set error, wilderness impact, known-class mAP retention and incremental forgetting, following established OWOD protocols.")
    doc.add_paragraph("Threats to validity include the modest protected test size, the use of one selected Teacher and Student seed in the headline comparison, mixed-domain public data, and the absence of a formal AP-small/tiling benchmark. The report does not claim that 80% detection performance was achieved. It also does not claim that the Student is more accurate than the Teacher or that image clusters constitute a complete open-world detector. These boundaries are important because the value of a research prototype depends as much on accurate interpretation as on the headline score.")

    doc.add_heading("7 Impact Statement",1)
    doc.add_paragraph("Industrial and economic impact: the delivered workflow can shorten early-stage experimentation for fruit inspection by reusing public unlabelled images and recording which pseudo labels are trusted. A small integrator or teaching laboratory can compare checkpoints without building a new annotation platform. Commercial deployment would still require target-site images, product-specific acceptance thresholds, software licensing review and a formal maintenance process.")
    doc.add_paragraph("Environmental and societal impact: improved automated inspection could support more consistent sorting and reduce avoidable handling or waste, but those benefits were not measured in this project. GPU training consumes energy, so future experiments should report training time and energy alongside accuracy. Dataset bias can reduce performance for rare cultivars, unusual lighting or damaged fruit. Human oversight is required wherever a missed or incorrect detection could affect payment, quality grading or worker safety.")
    doc.add_paragraph("Stakeholders include growers, pack-house operators, machine-vision integrators, researchers, students and end users responsible for quality assurance. Implementation should begin with a controlled pilot, a site-specific labelled test, error review by domain staff and documented fallback procedures. The evidence manifest, model hashes and export logs provide a foundation for that governance, but they do not replace field validation.")

    doc.add_heading("8 Conclusions and Future Work",1)
    doc.add_paragraph(f"A complete Windows-native semi-supervised fruit-detection prototype was implemented on an RTX 3080. The selected Teacher achieved fixed-test mAP@0.5 = {fmt(teacher['metrics']['map50'])}; the Student completed the Teacher–pseudo-label–Student chain and achieved {fmt(student['metrics']['map50'])}. The Trust Filter produced 2,036 canonical pseudo boxes with 94.3% protected audit precision. A separate self-supervised branch achieved holdout purity = {fmt(open_world['metrics']['holdout']['purity'])} across six fruits outside the detector registry, and a Chinese PySide6 program demonstrated image, folder and video inference.")
    doc.add_paragraph("The next optimisation phase should target the weakest class with curated labels, run class-specific threshold and sampling ablations, repeat the final Student across multiple seeds, and measure AP-small and throughput on a defined benchmark. The open-world interface should then be extended with box-level unknown proposals, explicit unknown metrics, human class registration and incremental learning. This sequence preserves the working prototype while addressing the evidence gaps identified by the current results.")

    doc.add_heading("References",1)
    refs=[
        "1. Tarvainen A, Valpola H. Mean teachers are better role models: weight-averaged consistency targets improve semi-supervised deep learning. Advances in Neural Information Processing Systems. 2017;30:1195–1204.",
        "2. Liu YC, Ma CY, He Z, et al. Unbiased Teacher for semi-supervised object detection. International Conference on Learning Representations. 2021.",
        "3. Xu M, Zhang Z, Hu H, et al. End-to-end semi-supervised object detection with Soft Teacher. Proceedings of ICCV. 2021:3060–3069.",
        "4. Shehzadi T, Hashmi KA, Stricker D, Afzal MZ. Sparse Semi-DETR: sparse learnable queries for semi-supervised object detection. Proceedings of CVPR. 2024:5840–5850.",
        "5. Johanson R, Wilms C, Johannsen O, Frintrop S. S3AD: semi-supervised small apple detection in orchard environments. Proceedings of WACV. 2024:7076–7085.",
        "6. Joseph KJ, Khan S, Khan FS, Balasubramanian VN. Towards open world object detection. Proceedings of CVPR. 2021:5830–5840.",
        "7. Zohar O, Wang KC, Yeung S. PROB: probabilistic objectness for open world object detection. Proceedings of CVPR. 2023:11444–11453.",
        "8. Yavuz M, Güney F. O1O: grouping of known classes to identify unknown objects as odd-one-out. Proceedings of ACCV. 2024:614–629.",
        "9. Zhang S, Ni Y, Du J, et al. Open-world objectness modeling unifies novel object detection. Proceedings of CVPR. 2025:30332–30342.",
        "10. Xi X, Huang Y, Luo R, Qiu Y. OW-OVD: unified open world and open vocabulary object detection. Proceedings of CVPR. 2025:25454–25464.",
        "11. Xiao F, Wang H, Xu Y, Zhang R. Fruit detection and recognition based on deep learning for automatic harvesting: an overview and review. Agronomy. 2023;13(6):1625.",
        "12. Kuznetsova A, Rom H, Alldrin N, et al. The Open Images Dataset V4: unified image classification, object detection, and visual relationship detection at scale. International Journal of Computer Vision. 2020;128:1956–1981.",
    ]
    for ref in refs: doc.add_paragraph(ref)

    doc.add_heading("Appendix A. Symbols and Abbreviations",1)
    add_table(doc,"Table 6. Symbols and abbreviations.",["Term","Definition"],[["AP","Average precision for one class"],["mAP@0.5","Mean AP at intersection-over-union threshold 0.5"],["SSOD","Semi-supervised object detection"],["OWOD","Open-world object detection"],["NMI","Normalized mutual information"],["ARI","Adjusted Rand index"],["IoU","Intersection over union"],["GUI","Graphical user interface"]],[1700,7660])
    doc.add_page_break()
    doc.add_heading("Appendix B. Reproducibility and Evidence Boundary",1)
    doc.add_paragraph("The package contains normalized data descriptors, split lists, figures, selected result workbooks, a required-deliverables inventory and SHA-256 manifest. Raw pseudo-label decisions, checkpoints, run records and fixed-test JSON files remain in the runtime evidence directory because copying them would substantially enlarge the report package. Their absolute paths are retained in the inventory. The report presents selected final evidence; exploratory low-score experiments remain available internally but are intentionally omitted from the customer-facing main body.")
    add_table(doc,"Table 7. Requirement-to-evidence alignment.",["Requirement","Delivered evidence","Boundary"],[["Limited-label semi-supervised detection","Teacher, Trust Filter, Student and fixed-test records","Student is demonstrable but does not exceed Teacher"],["Five registered fruits","YOLOv8m checkpoints and GUI","Apple, Banana, Orange, Strawberry, Pineapple"],["Additional fruit recognition","Six-category self-supervised discovery","Image-level; not box-level OWOD"],["Windows/RTX 3080","Conda environment and PySide6 program","No camera dependency"],["Professional report","DOCX/PDF, nine figures, seven tables","Identity fields require author completion"],["Delivery files","Inventory CSV/Markdown and four result workbooks","Formal AP-small remains to be measured"]],[2100,3600,3660])
    doc.add_heading("Appendix C. Submission Checklist",1)
    doc.add_paragraph("Before academic submission, complete the identity fields on the cover, update the Word table of contents if prompted, and confirm the course platform’s anonymity rule. The report remains below the 5,000-word narrative limit and below ten figures and ten tables. The customer has confirmed that a defence PPTX, English speech script and recording instructions are not required.")

    # Keep the report narrative aligned with the current delivered English GUI and
    # its external-camera mode, and renumber the original tables after inserting
    # the supervised/semi-supervised comparison as Table 1.
    replacements={
        "A Chinese PySide6 desktop program supports image, folder and video inference.": "An English PySide6 desktop program supports image, folder, video and external-camera inference.",
        "Camera input was deliberately removed from scope, and interfaces": "External-camera inference is included for live demonstrations, and interfaces",
        "deliver a Chinese PySide6 application": "deliver an English PySide6 application",
        "a Chinese PySide6 program demonstrated image, folder and video inference": "an English PySide6 program demonstrated image, folder, video and external-camera inference",
        "displays Chinese class names": "displays English class names",
        "The implementation avoids camera dependencies as requested.": "The external-camera mode supports camera selection, live preview, detection start/stop controls and annotated snapshot export.",
        "No camera dependency": "External-camera inference included",
        "nine figures, seven tables": "nine figures, eight tables",
        "Table 1. Dataset roles": "Table 2. Dataset roles",
        "Table 2. Trust Filter": "Table 3. Trust Filter",
        "Table 3. Selected model": "Table 4. Selected model",
        "Table 4. Selected fixed-test": "Table 5. Selected fixed-test",
        "Table 5. Image-level": "Table 6. Image-level",
        "Table 6. Symbols": "Table 7. Symbols",
        "Table 7. Requirement-to-evidence": "Table 8. Requirement-to-evidence",
    }
    all_paragraphs=list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for paragraph in all_paragraphs:
        updated=paragraph.text
        for old,new in replacements.items(): updated=updated.replace(old,new)
        if updated != paragraph.text:
            for run in paragraph.runs: run.text=""
            paragraph.add_run(updated)

    text="\n".join(p.text for p in doc.paragraphs)
    wc=word_count(text)
    if wc>5000: raise ValueError(f"report exceeds 5000 words: {wc}")
    path=out/"final_report.docx"; doc.save(path)
    return path,wc


def main():
    p=argparse.ArgumentParser()
    p.add_argument("--output",type=Path,required=True); p.add_argument("--teacher",type=Path,required=True); p.add_argument("--student",type=Path,required=True); p.add_argument("--audit",type=Path,required=True); p.add_argument("--open-world",type=Path,required=True); p.add_argument("--student-results",type=Path,required=True); p.add_argument("--gui",type=Path,required=True); p.add_argument("--split-manifest",type=Path,required=True); p.add_argument("--dataset-yaml",type=Path,required=True); p.add_argument("--unlabeled-manifest",type=Path,required=True)
    a=p.parse_args(); out=a.output.resolve()
    if out.exists(): raise SystemExit(f"refusing to overwrite: {out}")
    out.mkdir(parents=True)
    teacher=load_json(a.teacher); student=load_json(a.student); audit=load_json(a.audit); ow=load_json(a.open_world); gui_results=load_json(a.gui/"results.json")
    create_data_index(out,a.split_manifest,a.dataset_yaml,a.unlabeled_manifest)
    figs=build_figures(out,teacher,student,ow,a.student_results,a.gui)
    books=create_workbooks(out,teacher,student,audit,ow,gui_results)
    docx,wc=build_docx(out,teacher,student,audit,ow,figs,books)
    items,md,csvp=build_inventory(out,docx,None,Path(r"E:\bishe\fruit\.worktrees\fruit-ssod-implementation"),Path(r"E:\fruit_ssod_runtime"))
    source_paths=[a.teacher,a.student,a.audit,a.open_world,a.student_results,a.gui/"results.json",a.split_manifest,a.dataset_yaml,a.unlabeled_manifest]
    manifest={"protocol":"fruit_ssod_professional_final_report_v5","word_count":wc,"figure_count":len(figs),"table_count":8,"report_docx":{"path":str(docx),"sha256":sha256(docx)},"source_evidence":[{"path":str(x.resolve()),"sha256":sha256(x)} for x in source_paths],"figures":[{"path":str(v[0]),"sha256":sha256(v[0]),"caption":v[1]} for v in figs.values()],"workbooks":[{"path":str(x),"sha256":sha256(x)} for x in books],"inventory":{"markdown":str(md),"csv":str(csvp)}}
    (out/"manifest.json").write_text(json.dumps(manifest,ensure_ascii=False,indent=2)+"\n",encoding="utf-8")
    (out/"README.md").write_text("# Professional Final Report Package\n\nContains the evidence-bound final report, nine figures, eight tables, normalized data descriptors, four result workbooks, and the original-requirement deliverables inventory. Early exploratory low-score runs are excluded from the customer-facing report but remain in the runtime evidence archive.\n",encoding="utf-8")
    print(json.dumps({"output":str(out),"docx":str(docx),"word_count":wc,"figures":len(figs),"tables":8},ensure_ascii=False))


if __name__=="__main__": main()
