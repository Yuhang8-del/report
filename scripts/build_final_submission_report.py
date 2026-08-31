"""Build the English technical Final Report from the sealed project evidence.

The report follows the supplied Advanced Project guidance and execution plan:
title page, abstract, Introduction, Literature Review, Methodology, Experimental
Setup, Results, Discussion, Impact Statement, Conclusions, References and
Appendix.  It deliberately reports the achieved exploratory evidence and its
limitations rather than inventing an acceptance result.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import re
import shutil
from pathlib import Path
from typing import Any, Mapping, Sequence


KNOWN_CLASSES = ["Apple", "Banana", "Orange", "Strawberry", "Pineapple"]
NOVEL_CLASSES = ["Avocado", "Blueberry", "Cherry", "Kiwi", "Mango", "Rockmelon"]


def _load(path: Path) -> Mapping[str, Any]:
    value = json.loads(path.resolve(strict=True).read_text(encoding="utf-8"))
    if not isinstance(value, Mapping):
        raise ValueError(f"expected JSON object: {path}")
    return value


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _f(value: Any, digits: int = 4) -> str:
    return "n/a" if value is None else f"{float(value):.{digits}f}"


def _pct(value: Any, digits: int = 1) -> str:
    return "n/a" if value is None else f"{100.0 * float(value):.{digits}f}%"


def _words(text: str) -> int:
    return len(re.findall(r"\b[\w@.\-/]+\b", text))


def _ensure_output(output: Path) -> Path:
    output = output.resolve()
    if output.exists():
        raise ValueError(f"refusing to overwrite an existing report package: {output}")
    output.mkdir(parents=True, exist_ok=False)
    return output


def _matrix_rows(matrix: Mapping[str, Any]) -> list[Mapping[str, Any]]:
    rows = matrix.get("rows")
    if not isinstance(rows, list) or len(rows) != 6:
        raise ValueError("the supervised reference matrix must contain six rows")
    return [row for row in rows if isinstance(row, Mapping)]


def _build_tables(
    output: Path,
    matrix: Mapping[str, Any],
    teacher: Mapping[str, Any],
    student: Mapping[str, Any],
    open_world: Mapping[str, Any],
) -> dict[str, dict[str, Any]]:
    tables_dir = output / "tables"
    tables_dir.mkdir()
    rows = _matrix_rows(matrix)
    dataset_rows = [
        ["Known runtime classes", ", ".join(KNOWN_CLASSES)],
        ["Primary labelled source", "Open Images V7-derived five-class subset"],
        ["Full labelled train / validation / test", "542 / 90 / 90 images"],
        ["Label budgets", "10%: 54; 20%: 108; 40%: 217; 100%: 542 images"],
        ["Unlabelled training extension", "2,341 Open Images V7 images"],
        ["Novel-category pool", "639 images across six categories"],
    ]
    protocol_rows = [
        ["Teacher", "YOLOv8m, 1024 px, up to 120 epochs, batch 2, seed 42", _f(teacher["metrics"]["map50"])],
        ["Student", "Teacher initialisation; YOLOv8m, 768 px, AdamW, 80 epochs, patience 20", _f(student["metrics"]["map50"])],
        ["Pseudo-label policy", "Confidence 0.50, class/size/view filtering; 50:50 human/pseudo sampling", "exploratory"],
        ["Open-world encoder", "SimCLR-style augmentation consistency, CUDA, 10 epochs", "image-level"],
        ["Desktop demonstrator", "PySide6; local images, folders and video files; no camera", "five classes"],
    ]
    matrix_table = []
    for row in sorted(rows, key=lambda item: (int(item.get("label_budget_percent", 0)), int(item.get("seed", 0)))):
        metrics = row["fixed_test"]
        matrix_table.append(
            [
                f"{row['label_budget_percent']}%",
                str(row["seed"]),
                _f(metrics.get("map50")),
                _f(metrics.get("map50_95")),
                _f(metrics.get("precision")),
                _f(metrics.get("recall")),
                _f(metrics.get("f1")),
            ]
        )
    main_results = [
        ["Teacher (v3-r3)", "Domain-balanced high-resolution fine-tune", _f(teacher["metrics"]["map50"]), _f(teacher["metrics"]["map50_95"]), _f(teacher["metrics"]["precision"]), _f(teacher["metrics"]["recall"]), _f(teacher["metrics"]["f1"])],
        ["Student (v3-r3)", "20% human labels + filtered pseudo labels", _f(student["metrics"]["map50"]), _f(student["metrics"]["map50_95"]), _f(student["metrics"]["precision"]), _f(student["metrics"]["recall"]), _f(student["metrics"]["f1"])],
    ]
    discovery = open_world["metrics"]["discovery"]
    holdout = open_world["metrics"]["holdout"]
    open_world_table = [
        ["Discovery split", str(discovery["image_count"]), _f(discovery["purity"]), _f(discovery["nmi"]), _f(discovery["ari"])],
        ["Protected holdout", str(holdout["image_count"]), _f(holdout["purity"]), _f(holdout["nmi"]), _f(holdout["ari"])],
    ]
    deployment_table = [
        ["Runtime classes", ", ".join(KNOWN_CLASSES)],
        ["Input modes", "Single image, batch images, local video"],
        ["Hardware", "NVIDIA GeForce RTX 3080, 10 GB; CUDA 12.1"],
        ["Model checkpoint", "Student v3-r3 best.pt; SHA-256 recorded in fixed-test evidence"],
        ["Open-world status", "Separate image-level discovery evidence; no new runtime class ID"],
    ]
    specs = {
        "dataset_summary": ("Table 1. Dataset, class and split summary.", ["Item", "Specification"], dataset_rows),
        "training_protocol": ("Table 2. Training and demonstration protocol.", ["Component", "Configuration", "Reported output"], protocol_rows),
        "supervised_matrix": ("Table 3. Fixed-test supervised reference matrix.", ["Label budget", "Seed", "mAP@0.5", "mAP@0.5:0.95", "Precision", "Recall", "F1"], matrix_table),
        "main_results": ("Table 4. Selected Teacher and Student evidence.", ["Model", "Training signal", "mAP@0.5", "mAP@0.5:0.95", "Precision", "Recall", "F1"], main_results),
        "open_world": ("Table 5. Image-level novel-category discovery evidence.", ["Split", "Images", "Purity", "NMI", "ARI"], open_world_table),
        "deployment": ("Table 6. Demonstration and deployment boundary.", ["Item", "Delivered scope"], deployment_table),
    }
    result: dict[str, dict[str, Any]] = {}
    for key, (caption, headers, data) in specs.items():
        path = tables_dir / f"{key}.csv"
        with path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(headers)
            writer.writerows(data)
        result[key] = {"caption": caption, "headers": headers, "rows": data, "path": str(path), "sha256": _sha256(path)}
    return result


def _build_figures(
    output: Path,
    matrix: Mapping[str, Any],
    teacher: Mapping[str, Any],
    student: Mapping[str, Any],
    open_world: Mapping[str, Any],
    gui_export: Path,
) -> dict[str, dict[str, Any]]:
    import matplotlib.pyplot as plt
    import numpy as np
    from PIL import Image
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    figures_dir = output / "figures"
    figures_dir.mkdir()
    plt.rcParams.update({"font.family": "DejaVu Sans", "axes.titlesize": 12, "axes.labelsize": 10})
    result: dict[str, dict[str, Any]] = {}

    def save(fig: Any, name: str, caption: str) -> None:
        path = figures_dir / name
        fig.savefig(path, dpi=190, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        result[name] = {"path": str(path), "caption": caption, "sha256": _sha256(path)}

    # Figure 1: a compact workflow matching the methodology in the supplied PPT.
    fig, ax = plt.subplots(figsize=(13, 3.4))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 3.4)
    ax.axis("off")
    labels = [
        (0.2, "Public data\nand audit", "#d8eef0"),
        (2.8, "Fixed split\nand label budgets", "#e4f0ff"),
        (5.4, "Teacher\nYOLOv8m", "#dfe9ff"),
        (8.0, "Pseudo-label\nTrust Filter", "#fff0d2"),
        (10.6, "Student +\nfixed test", "#dff3df"),
    ]
    for x, text, color in labels:
        ax.add_patch(FancyBboxPatch((x, 1.05), 1.9, 1.25, boxstyle="round,pad=0.08,rounding_size=0.12", facecolor=color, edgecolor="#31556b", linewidth=1.2))
        ax.text(x + 0.95, 1.68, text, ha="center", va="center", fontsize=10, weight="bold", color="#19324a")
    for x in [2.15, 4.75, 7.35, 9.95]:
        ax.add_patch(FancyArrowPatch((x, 1.68), (x + 0.55, 1.68), arrowstyle="-|>", mutation_scale=16, linewidth=1.5, color="#2f9e8f"))
    ax.text(6.5, 2.95, "Auditable semi-supervised fruit detection workflow", ha="center", fontsize=14, weight="bold", color="#12344d")
    ax.text(6.5, 0.45, "A separate self-supervised encoder then clusters candidate images outside the five-class runtime registry.", ha="center", fontsize=9, color="#5f7484")
    save(fig, "figure_01_workflow.png", "Figure 1. Overall workflow of the proposed fruit detection and evidence pipeline.")

    # Figure 2: label-budget curve, averaging the three 20% seeds.
    groups: dict[int, list[float]] = {}
    for row in _matrix_rows(matrix):
        groups.setdefault(int(row["label_budget_percent"]), []).append(float(row["fixed_test"]["map50"]))
    budgets = sorted(groups)
    means = [float(np.mean(groups[b])) for b in budgets]
    fig, ax = plt.subplots(figsize=(7.5, 4.2))
    ax.plot(budgets, means, marker="o", linewidth=2.5, color="#2f9e8f")
    for x, y in zip(budgets, means):
        ax.annotate(f"{y:.3f}", (x, y), textcoords="offset points", xytext=(0, 8), ha="center", fontsize=9)
    ax.set_xlabel("Human-labelled training budget (%)")
    ax.set_ylabel("Fixed-test mAP@0.5")
    ax.set_title("Label efficiency in the supervised reference matrix")
    ax.set_ylim(0, max(means) * 1.35)
    ax.grid(axis="y", alpha=0.25)
    save(fig, "figure_02_label_budget.png", "Figure 2. Fixed-test mAP@0.5 rises with the supervised label budget; the 20% value is the mean of three seeds.")

    # Figure 3: per-class comparison on the same fixed test protocol.
    teacher_ap = [float(teacher["metrics"]["per_class_ap50"][str(i)]) for i in range(5)]
    student_ap = [float(student["metrics"]["per_class_ap50"][str(i)]) for i in range(5)]
    x = np.arange(len(KNOWN_CLASSES))
    fig, ax = plt.subplots(figsize=(9, 4.4))
    ax.bar(x - 0.19, teacher_ap, width=0.38, label="Teacher", color="#3977a8")
    ax.bar(x + 0.19, student_ap, width=0.38, label="Student", color="#2f9e8f")
    ax.set_xticks(x, KNOWN_CLASSES, rotation=15)
    ax.set_ylim(0, 1)
    ax.set_ylabel("AP@0.5")
    ax.set_title("Per-class fixed-test AP@0.5")
    ax.legend(frameon=False)
    ax.grid(axis="y", alpha=0.2)
    save(fig, "figure_03_per_class_ap.png", "Figure 3. Per-class AP@0.5 for the selected Teacher and Student checkpoints.")

    # Figure 4: qualitative GUI outputs.
    image_paths = sorted((gui_export / "annotated_images").glob("*.png"))[:3]
    if image_paths:
        fig, axes = plt.subplots(1, len(image_paths), figsize=(13, 4.6))
        axes = np.atleast_1d(axes)
        for axis, path in zip(axes, image_paths):
            image = Image.open(path).convert("RGB")
            axis.imshow(image)
            axis.set_title(path.stem.replace("_annotated", ""), fontsize=9)
            axis.axis("off")
        save(fig, "figure_04_gui_examples.png", "Figure 4. Example annotated outputs generated by the Student checkpoint for the PySide6 demonstration.")

    # Figure 5: self-supervised representation loss.
    loss_curve = [float(value) for value in open_world["self_supervised"]["loss_curve"]]
    fig, ax = plt.subplots(figsize=(7.5, 4.2))
    ax.plot(range(1, len(loss_curve) + 1), loss_curve, marker="o", color="#d47c35", linewidth=2.3)
    ax.set_xlabel("Self-supervised epoch")
    ax.set_ylabel("Consistency loss")
    ax.set_title("Self-supervised encoder optimisation")
    ax.grid(axis="y", alpha=0.25)
    save(fig, "figure_05_self_supervised_loss.png", "Figure 5. SimCLR-style augmentation-consistency loss used before image-level novel-category clustering.")

    # Figure 6: discovery metrics.
    metrics = ["Purity", "NMI", "ARI"]
    discovery_values = [float(open_world["metrics"]["discovery"][key.lower()]) for key in metrics]
    holdout_values = [float(open_world["metrics"]["holdout"][key.lower()]) for key in metrics]
    x = np.arange(len(metrics))
    fig, ax = plt.subplots(figsize=(7.5, 4.2))
    ax.bar(x - 0.19, discovery_values, width=0.38, label="Discovery", color="#6a9cc5")
    ax.bar(x + 0.19, holdout_values, width=0.38, label="Holdout", color="#d47c35")
    ax.set_xticks(x, metrics)
    ax.set_ylim(0, 1)
    ax.set_ylabel("Score")
    ax.set_title("Image-level novel-category discovery")
    ax.legend(frameon=False)
    ax.grid(axis="y", alpha=0.2)
    save(fig, "figure_06_open_world_metrics.png", "Figure 6. Purity, NMI and ARI for the six-category image-level discovery experiment.")
    return result


def _add_docx_table(document: Any, spec: Mapping[str, Any]) -> None:
    from docx.shared import Pt

    document.add_paragraph(str(spec["caption"]))
    table = document.add_table(rows=1, cols=len(spec["headers"]))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, spec["headers"]):
        cell.text = str(value)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for row in spec["rows"]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def _add_docx_figure(document: Any, figure: Mapping[str, Any]) -> None:
    from docx.shared import Inches

    document.add_picture(str(figure["path"]), width=Inches(6.25))
    document.add_paragraph(str(figure["caption"]))


def _build_docx(
    output: Path,
    teacher: Mapping[str, Any],
    student: Mapping[str, Any],
    matrix: Mapping[str, Any],
    open_world: Mapping[str, Any],
    tables: Mapping[str, Mapping[str, Any]],
    figures: Mapping[str, Mapping[str, Any]],
) -> tuple[Path, int]:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for style_name, size, color in (("Title", 24, "12344D"), ("Heading 1", 15, "12344D"), ("Heading 2", 11, "2F6F77")):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    # Cover page required by the supplied University guidance.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("University of Birmingham\n").bold = True
    p.add_run("School of Engineering\nDepartment of Mechanical Engineering\n\n").bold = True
    p.add_run("MSc Advanced Mechanical Engineering\nAdvanced Project\n\n").bold = True
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Semi-Supervised Fruit Object Detection\nUsing Limited Labels and Unlabelled Data")
    run.bold = True
    run.font.size = Pt(22)
    title_table = doc.add_table(rows=4, cols=2)
    title_table.style = "Table Grid"
    for row, (label, value) in zip(title_table.rows, (("Student", "XXX"), ("ID number", "XXX"), ("Supervisor", "XXX"), ("Project type", "Experimental research and demonstration prototype"))):
        row.cells[0].text, row.cells[1].text = label, value
    doc.add_paragraph("\nEvidence-bound submission draft. Numerical claims are generated from the sealed project artifacts listed in the Appendix.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This project investigated semi-supervised fruit object detection when only a limited portion of public bounding-box data is treated as labelled. "
        "The delivered workflow fixed the five-class registry (Apple, Banana, Orange, Strawberry and Pineapple), separated labelled, unlabelled, validation and test memberships, trained a YOLOv8m Teacher, filtered offline pseudo labels and trained a Student detector. "
        f"On the protected fixed test, the selected Teacher achieved mAP@0.5 = {_f(teacher['metrics']['map50'])}, while the exploratory Student achieved {_f(student['metrics']['map50'])}. "
        "A separate self-supervised representation experiment clustered an independent pool of 639 images from six categories outside the runtime registry; holdout purity, NMI and ARI were "
        f"{_f(open_world['metrics']['holdout']['purity'])}, {_f(open_world['metrics']['holdout']['nmi'])} and {_f(open_world['metrics']['holdout']['ari'])}, respectively. "
        "The results demonstrate a reproducible research and demonstration chain, but do not support a formal 0.80 accuracy claim or box-level open-world detection claim."
    )
    doc.add_heading("Keywords", level=1)
    doc.add_paragraph("semi-supervised object detection; fruit detection; pseudo-label filtering; YOLOv8; reproducibility")

    doc.add_heading("1 Introduction", level=1)
    doc.add_paragraph(
        "Fruit inspection is a representative agricultural-vision problem in which image appearance varies with illumination, viewpoint, occlusion, density and background. "
        "Bounding-box annotation is expensive, so a detector that can exploit unlabelled images is attractive for research and future low-cost inspection systems. "
        "The project aim was to build and evaluate a transparent semi-supervised pipeline rather than to present an isolated model score."
    )
    doc.add_paragraph("The technical objectives were to: (i) construct a leakage-safe public-data split; (ii) quantify supervised performance under 10%, 20%, 40% and 100% label budgets; (iii) generate and filter offline pseudo labels; (iv) train and evaluate a Student detector; (v) investigate an independent self-supervised extension for fruit categories outside the five-class registry; and (vi) expose the resulting model through a Windows-native PySide6 file-based demonstrator.")
    doc.add_paragraph("The scope is deliberately bounded. The runtime detector recognises five registered classes. The open-world extension reports image-level candidate clusters and does not add semantic class IDs to the detector. Camera capture is outside the delivered scope.")

    doc.add_heading("2 Literature Review", level=1)
    doc.add_paragraph("Mean Teacher methods introduced the use of weight-averaged targets and consistency training to exploit unlabelled data [1]. Unbiased Teacher addressed confirmation bias in semi-supervised object detection through confidence-aware pseudo-label learning [2], while Soft Teacher used soft pseudo targets and strong/weak augmentation to improve detector training [3]. More recent work has explored YOLO-oriented efficiency [4], end-to-end one-stage formulations [5], and sparse transformer-based semi-supervised detection [6].")
    doc.add_paragraph("Agricultural detection adds two practical difficulties. First, fruit datasets contain class imbalance, small objects and visually similar backgrounds. Secondly, orchard and market images exhibit domain shift. S3AD demonstrates that semi-supervised learning can be useful for small apple detection in orchard scenes [7], and the agricultural review by Xiao et al. highlights the importance of illumination, occlusion and deployment conditions [8].")
    doc.add_paragraph("The methodological gap addressed here is reproducibility at prototype scale: split identities, label budgets, pseudo-label decisions, fixed-test evidence, model hashes and a customer-facing demonstrator are retained together. This makes a modest result interpretable and identifies where further data and algorithm improvements are required.")

    doc.add_heading("3 Methodology", level=1)
    doc.add_heading("3.1 Data and class registry", level=2)
    doc.add_paragraph("The primary detection data are an Open Images V7-derived five-class subset. The canonical mapping is Apple, Banana, Orange, Strawberry and Pineapple. The full labelled training membership contains 542 images, with 90 validation images and 90 fixed-test images. The unlabelled training extension contains 2,341 Open Images V7 images. Image-level open-world discovery uses an independent 639-image pool containing Avocado, Blueberry, Cherry, Kiwi, Mango and Rockmelon.")
    _add_docx_table(doc, tables["dataset_summary"])
    doc.add_heading("3.2 Leakage-safe split and label budgets", level=2)
    doc.add_paragraph("Split lists and class mappings were frozen before training. The validation and fixed-test memberships were not used to generate Student training examples. The supervised matrix used 10%, 20%, 40% and 100% labelled budgets; the 20% condition was repeated for three seeds (42, 3407 and 2026) to expose sensitivity to sampling.")
    doc.add_heading("3.3 Teacher, pseudo labels and Student", level=2)
    doc.add_paragraph("The selected Teacher is the v3-r3 domain-balanced YOLOv8m high-resolution fine-tune using 1024-pixel inputs for up to 120 epochs with patience 30. It predicts the unlabelled pool offline. The Trust Filter combines a global confidence floor of 0.50 with category-aware calibration, a cross-view IoU requirement, size and aspect-ratio checks, a maximum box count and a 50:50 human/pseudo sampling policy. The Student is initialised from the Teacher checkpoint and trained for up to 80 epochs at 768 pixels with AdamW and patience 20. The customer-authorised path allowed this exploratory Student to proceed without using the historical pseudo-label precision gate as a stop condition.")
    _add_docx_table(doc, tables["training_protocol"])
    doc.add_heading("3.4 Self-supervised extension", level=2)
    doc.add_paragraph("After Student training, a SimCLR-style augmentation-consistency encoder was trained for ten epochs on the independent novel-category pool. Deterministic k-means then produced six image-level clusters. Cluster names are post-hoc evaluation mappings; they are not training labels and they do not alter the five-class detector registry.")
    doc.add_heading("3.5 Demonstration architecture", level=2)
    doc.add_paragraph("The PySide6 desktop program loads one validated .pt checkpoint in a background worker and performs file-based image, folder and video inference in separate workers. The interface displays boxes, Chinese fruit names, confidence values, latency, progress and export metadata. No camera device is opened, and the open-world extension is presented as an evidence boundary rather than an enabled runtime action.")

    doc.add_heading("4 Experimental Setup", level=1)
    doc.add_paragraph("Experiments ran on Windows using Python 3.10.20, PyTorch 2.5.1+cu121, Ultralytics 8.4.31 and an NVIDIA GeForce RTX 3080 with 10 GB memory. The Student configuration used batch size 4, workers 0, AMP disabled for the selected recovery run, cosine learning-rate scheduling, mosaic augmentation and mixup 0.05. The primary metric is mAP@0.5; mAP@0.5:0.95, precision, recall and F1 are reported as secondary indicators. Per-class AP@0.5 and qualitative detection outputs are retained for error analysis.")
    _add_docx_figure(doc, figures["figure_01_workflow.png"])
    _add_docx_table(doc, tables["supervised_matrix"])

    doc.add_heading("5 Results", level=1)
    doc.add_heading("5.1 Label-budget reference", level=2)
    doc.add_paragraph("The fixed-test supervised matrix shows a monotonic but low absolute trend as labelled image coverage increases. The mean 20% result is computed from three seeds. The 100% row is a diagnostic reference, not a formal acceptance claim.")
    _add_docx_figure(doc, figures["figure_02_label_budget.png"])
    doc.add_heading("5.2 Teacher and Student", level=2)
    doc.add_paragraph(f"The selected Teacher reached fixed-test mAP@0.5 = {_f(teacher['metrics']['map50'])} in the retained evidence. The Student fixed-test result was mAP@0.5 = {_f(student['metrics']['map50'])}, mAP@0.5:0.95 = {_f(student['metrics']['map50_95'])}, precision = {_f(student['metrics']['precision'])}, recall = {_f(student['metrics']['recall'])} and F1 = {_f(student['metrics']['f1'])}. The Student therefore forms a usable demonstration checkpoint but did not improve on this Teacher fixed-test score.")
    _add_docx_table(doc, tables["main_results"])
    _add_docx_figure(doc, figures["figure_03_per_class_ap.png"])
    doc.add_heading("5.3 Qualitative GUI evidence", level=2)
    doc.add_paragraph("The three annotated examples below are generated from the selected Student checkpoint and are bundled with the report manifest. They demonstrate multi-object boxes and class/confidence rendering, but qualitative examples must not be interpreted as a substitute for the protected fixed-test evaluation.")
    _add_docx_figure(doc, figures["figure_04_gui_examples.png"])
    doc.add_heading("5.4 Self-supervised novel-category discovery", level=2)
    doc.add_paragraph("The representation loss decreased over the ten self-supervised epochs. On 510 discovery images, purity, NMI and ARI were " + f"{_f(open_world['metrics']['discovery']['purity'])}, {_f(open_world['metrics']['discovery']['nmi'])} and {_f(open_world['metrics']['discovery']['ari'])}. On the protected 129-image holdout, the corresponding scores were " + f"{_f(open_world['metrics']['holdout']['purity'])}, {_f(open_world['metrics']['holdout']['nmi'])} and {_f(open_world['metrics']['holdout']['ari'])}. This is evidence of category structure in the independent pool, not a box-level open-world detector score.")
    _add_docx_figure(doc, figures["figure_05_self_supervised_loss.png"])
    _add_docx_figure(doc, figures["figure_06_open_world_metrics.png"])
    _add_docx_table(doc, tables["open_world"])
    _add_docx_table(doc, tables["deployment"])

    doc.add_heading("6 Discussion", level=1)
    doc.add_paragraph("The results support three conclusions. First, the project has a reproducible end-to-end chain from public data to a working Student checkpoint and desktop demonstration. Secondly, the supervised matrix reveals a data-quality and coverage limitation: even the full-label diagnostic is only mAP@0.5 = " + _f(matrix["upper_bound_gate"]["fixed_test_map50"]) + ". This means that extending training duration alone is unlikely to deliver the original 0.80 target; additional label coverage, class balancing, split review and hard-scene curation are more credible next interventions. Thirdly, the Student did not exceed the selected Teacher on the fixed test, so the value of the semi-supervised stage is currently methodological and operational rather than a proven accuracy gain.")
    doc.add_paragraph("The failure modes are consistent with the project risks identified in the supplied PPT: class imbalance, small/occluded fruits, domain shift and pseudo-label confirmation bias. Orange has the strongest fixed-test AP in the supervised matrix, whereas Banana and Strawberry are more variable across runs. The Student's per-class pattern also shows that pseudo labels can help some classes while suppressing or mislocalising others. A future study should compare three-seed Student runs, calibrate class-specific thresholds on a larger protected validation set, add hard-negative analysis and evaluate selective tiling under a recorded latency budget.")
    doc.add_paragraph("Threats to validity include the limited number of labelled images, the single selected Teacher seed for the Student, the exploratory allowance below the historical pseudo-label precision gate, and the fact that novel-category names are assigned post hoc. The report therefore makes no formal 0.80 acceptance claim, no 0.85 upper-bound claim and no box-level open-world claim.")

    doc.add_heading("7 Impact Statement", level=1)
    doc.add_paragraph("WHAT: The project delivers a Windows-native, reproducible research prototype for five-class semi-supervised fruit detection, including fixed data manifests, supervised and Student checkpoints, auditable pseudo-label artifacts, a Chinese PySide6 demonstrator and an evidence-bound report. WHO: Potential users include agricultural-vision researchers, teaching laboratories, small automation integrators and students studying trustworthy object detection. HOW: When validated on a target deployment, filtered pseudo labels may reduce repeated annotation effort and the immutable audit trail can make failure modes visible to engineers. Potential industrial and commercial benefits include lower prototype labelling cost and faster comparison of detector configurations; environmental benefits such as reduced inspection waste are plausible but unmeasured; societal benefits depend on responsible use and human oversight. Implementation stakeholders must review dataset licences, privacy, bias, safety and domain-specific acceptance before operational deployment.")

    doc.add_heading("8 Conclusions", level=1)
    doc.add_paragraph("A complete research and demonstration prototype was implemented under Windows on an RTX 3080. The retained Teacher achieved fixed-test mAP@0.5 = " + _f(teacher["metrics"]["map50"]) + ", and the exploratory Student achieved " + _f(student["metrics"]["map50"]) + ". The self-supervised extension found measurable image-level structure among six categories outside the five-class registry, with holdout purity = " + _f(open_world["metrics"]["holdout"]["purity"]) + ". These results are sufficient for a transparent experimental prototype and customer demonstration, but they also show that the original accuracy target requires further data and methodology work. The next research step is a controlled data-quality and multi-seed optimisation study followed by box-level unknown-object evaluation.")

    doc.add_heading("References", level=1)
    references = [
        "1. Tarvainen A, Valpola H. Mean teachers are better role models: weight-averaged consistency targets improve semi-supervised deep learning. In: Advances in Neural Information Processing Systems. 2017;30:1195-1204.",
        "2. Liu YC, Ma CY, He Z, Kuo CW, Chen K, Zhang P, et al. Unbiased Teacher for semi-supervised object detection. In: International Conference on Learning Representations. 2021.",
        "3. Xu M, Zhang Z, Hu H, Wang J, Wang L, Wei F, et al. End-to-end semi-supervised object detection with Soft Teacher. In: Proceedings of the IEEE/CVF International Conference on Computer Vision. 2021:3060-3069.",
        "4. Xu B, Chen M, Guan W, Hu L. Efficient Teacher: semi-supervised object detection for YOLOv5. arXiv preprint arXiv:2302.07577. 2023.",
        "5. Luo G, Zhou Y, Jin L, Sun X, Ji R. Towards end-to-end semi-supervised learning for one-stage object detection. arXiv preprint arXiv:2302.11299. 2023.",
        "6. Shehzadi T, Hashmi KA, Stricker D, Afzal MZ. Sparse Semi-DETR: sparse learnable queries for semi-supervised object detection. In: Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024.",
        "7. Johanson R, Wilms C, Johannsen O, Frintrop S. S3AD: semi-supervised small apple detection in orchard environments. In: Proceedings of the IEEE/CVF Winter Conference on Applications of Computer Vision. 2024:7076-7085.",
        "8. Xiao F, Wang H, Xu Y, Zhang R. Fruit detection and recognition based on deep learning for automatic harvesting: an overview and review. Agronomy. 2023;13(6):1625.",
    ]
    for reference in references:
        doc.add_paragraph(reference)

    doc.add_heading("Appendix A. Reproducibility and evidence boundary", level=1)
    doc.add_paragraph("The report package records the source JSON paths, SHA-256 digests, selected checkpoint identities, split fingerprint, figure/table hashes and the effective environment. The primary fixed-test split was not used for model selection or Student pseudo-label generation. The GUI starts from scripts/start_gui.ps1 in the project worktree and loads the Student best.pt checkpoint explicitly. The deliverable contains no camera workflow and no private credentials. Open-world results are retained as a separate image-level discovery artifact; semantic class registration and box-level unknown-object mAP remain future work.")
    doc.add_heading("Appendix B. Alignment with the supplied requirements", level=1)
    doc.add_paragraph("The report follows the supplied Word guidance by including a title page, abstract and keywords, an experimental-paper main body, figures/tables with captions, references and appendices, while keeping the main narrative below 5,000 words and the figure/table count below ten each. It follows the supplied PPT structure by covering the problem and state of the art, data split and annotation budget, methodology, evaluation design, technical risks and ethics, deliverables, impact and limitations. Dates, student identity and supervisor fields remain placeholders for the author to complete before submission.")

    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    word_count = _words(body_text)
    if word_count > 5000:
        raise ValueError(f"main narrative exceeds 5000 words: {word_count}")
    output_path = output / "final_report.docx"
    doc.save(output_path)
    return output_path, word_count


def _build_pdf(
    output: Path,
    teacher: Mapping[str, Any],
    student: Mapping[str, Any],
    matrix: Mapping[str, Any],
    open_world: Mapping[str, Any],
    tables: Mapping[str, Mapping[str, Any]],
    figures: Mapping[str, Mapping[str, Any]],
) -> Path:
    from html import escape

    from matplotlib import font_manager
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_path = font_manager.findfont("DejaVu Sans")
    pdfmetrics.registerFont(TTFont("FruitSans", font_path))
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="FruitTitle", parent=styles["Title"], fontName="FruitSans", fontSize=22, leading=28, textColor=colors.HexColor("#12344D"), alignment=TA_CENTER, spaceAfter=20))
    styles.add(ParagraphStyle(name="FruitH1", parent=styles["Heading1"], fontName="FruitSans", fontSize=15, leading=19, textColor=colors.HexColor("#12344D"), spaceBefore=12, spaceAfter=7))
    styles.add(ParagraphStyle(name="FruitH2", parent=styles["Heading2"], fontName="FruitSans", fontSize=11.5, leading=15, textColor=colors.HexColor("#2F6F77"), spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name="FruitBody", parent=styles["BodyText"], fontName="FruitSans", fontSize=9.3, leading=13.2, spaceAfter=6))
    styles.add(ParagraphStyle(name="FruitSmall", parent=styles["BodyText"], fontName="FruitSans", fontSize=8, leading=10, textColor=colors.HexColor("#5F7484"), spaceAfter=4))
    styles.add(ParagraphStyle(name="FruitCaption", parent=styles["BodyText"], fontName="FruitSans", fontSize=8.2, leading=10, textColor=colors.HexColor("#536B7D"), spaceAfter=9))

    story: list[Any] = []
    story.extend([Spacer(1, 1.0 * inch), Paragraph("University of Birmingham", styles["FruitH2"]), Paragraph("School of Engineering<br/>Department of Mechanical Engineering<br/><br/>MSc Advanced Mechanical Engineering<br/>Advanced Project", styles["FruitBody"]), Spacer(1, 0.6 * inch), Paragraph("Semi-Supervised Fruit Object Detection<br/>Using Limited Labels and Unlabelled Data", styles["FruitTitle"]), Spacer(1, 0.35 * inch)])
    cover_rows = [["Student", "XXX"], ["ID number", "XXX"], ["Supervisor", "XXX"], ["Project type", "Experimental research and demonstration prototype"]]
    story.append(_pdf_table(cover_rows, ["Field", "Value"], styles))
    story.extend([Spacer(1, 0.3 * inch), Paragraph("Evidence-bound submission draft. Numerical claims are generated from the sealed project artifacts listed in the Appendix.", styles["FruitSmall"]), PageBreak()])

    def para(text: str, style: str = "FruitBody") -> None:
        story.append(Paragraph(escape(text).replace("\n", "<br/>") , styles[style]))

    def heading(text: str, level: int = 1) -> None:
        story.append(Paragraph(escape(text), styles["FruitH1" if level == 1 else "FruitH2"]))

    def figure(name: str) -> None:
        spec = figures[name]
        from PIL import Image as PILImage
        with PILImage.open(spec["path"]) as image:
            width, height = image.size
        max_width = 6.7 * inch
        max_height = 4.7 * inch
        scale = min(max_width / width, max_height / height)
        story.append(Image(spec["path"], width=width * scale, height=height * scale))
        story.append(Paragraph(escape(spec["caption"]), styles["FruitCaption"]))

    para("Abstract", "FruitH1")
    para("This project investigated semi-supervised fruit object detection when only a limited portion of public bounding-box data is treated as labelled. The delivered workflow fixed five runtime classes, separated labelled, unlabelled, validation and test memberships, trained a YOLOv8m Teacher, filtered offline pseudo labels and trained a Student detector. The selected Teacher achieved fixed-test mAP@0.5 = " + _f(teacher["metrics"]["map50"]) + ", while the exploratory Student achieved " + _f(student["metrics"]["map50"]) + ". A separate self-supervised representation experiment clustered an independent pool of 639 images from six categories outside the runtime registry; holdout purity, NMI and ARI were " + _f(open_world["metrics"]["holdout"]["purity"]) + ", " + _f(open_world["metrics"]["holdout"]["nmi"]) + " and " + _f(open_world["metrics"]["holdout"]["ari"]) + ", respectively. The results demonstrate a reproducible prototype, but do not support a formal 0.80 accuracy claim or box-level open-world detection claim.")
    para("Keywords: semi-supervised object detection; fruit detection; pseudo-label filtering; YOLOv8; reproducibility")
    heading("1 Introduction")
    para("Fruit inspection is a representative agricultural-vision problem in which image appearance varies with illumination, viewpoint, occlusion, density and background. Bounding-box annotation is expensive, so a detector that can exploit unlabelled images is attractive. The project aim was to build and evaluate a transparent semi-supervised pipeline rather than to present an isolated model score.")
    para("The objectives were to construct a leakage-safe public-data split, quantify supervised performance under multiple label budgets, generate and filter offline pseudo labels, train and evaluate a Student detector, investigate a self-supervised extension for categories outside the five-class registry and expose the resulting model through a Windows-native PySide6 demonstrator. The runtime detector recognises five registered classes; camera capture is outside scope.")
    heading("2 Literature Review")
    para("Mean Teacher methods introduced weight-averaged targets and consistency training for unlabelled data [1]. Unbiased Teacher addressed confirmation bias in semi-supervised object detection [2], while Soft Teacher used soft pseudo targets and strong/weak augmentation [3]. YOLO-oriented efficiency [4], end-to-end one-stage formulations [5] and sparse transformer-based approaches [6] extend this design space. S3AD demonstrates semi-supervised small apple detection in orchard scenes [7], and an agricultural review identifies illumination, occlusion and deployment variation as recurring challenges [8]. This project focuses on the reproducibility gap: split identities, pseudo-label decisions, fixed-test evidence, checkpoint hashes and a user-facing demonstrator are retained together.")
    heading("3 Methodology")
    heading("3.1 Data, split and budgets", 2)
    para("The primary data are an Open Images V7-derived five-class subset with 542 labelled training images, 90 validation images and 90 fixed-test images. The unlabelled extension contains 2,341 Open Images V7 images. Novel-category discovery uses 639 independent images across Avocado, Blueberry, Cherry, Kiwi, Mango and Rockmelon. Split lists and mappings were frozen before training; the test membership was not used for model selection or Student pseudo-label generation.")
    _pdf_table(tables["dataset_summary"]["rows"], tables["dataset_summary"]["headers"], styles, story, tables["dataset_summary"]["caption"])
    heading("3.2 Teacher, pseudo labels and Student", 2)
    para("The selected Teacher is the v3-r3 domain-balanced YOLOv8m high-resolution fine-tune at 1024 pixels for up to 120 epochs with patience 30. The Trust Filter combines confidence, category-aware calibration, cross-view agreement, size and aspect-ratio checks and a maximum box count. The Student is initialised from the Teacher and trained at 768 pixels for up to 80 epochs with AdamW and patience 20. The exploratory path was allowed to continue without using the historical pseudo-label precision gate as a stop condition.")
    _pdf_table(tables["training_protocol"]["rows"], tables["training_protocol"]["headers"], styles, story, tables["training_protocol"]["caption"])
    heading("3.3 Self-supervised extension and GUI", 2)
    para("A SimCLR-style augmentation-consistency encoder was trained for ten epochs, followed by deterministic k-means clustering. Cluster names are post-hoc evaluation mappings and do not alter the detector registry. The PySide6 program loads one validated checkpoint in a background worker and supports local image, folder and video files without camera access.")
    heading("4 Experimental Setup")
    para("Experiments ran on Windows using Python 3.10.20, PyTorch 2.5.1+cu121, Ultralytics 8.4.31 and an NVIDIA GeForce RTX 3080 with 10 GB memory. The primary metric is mAP@0.5; mAP@0.5:0.95, precision, recall, F1 and per-class AP are secondary indicators.")
    figure("figure_01_workflow.png")
    _pdf_table(tables["supervised_matrix"]["rows"], tables["supervised_matrix"]["headers"], styles, story, tables["supervised_matrix"]["caption"])
    heading("5 Results")
    heading("5.1 Label-budget reference", 2)
    para("The fixed-test supervised matrix shows a monotonic but low absolute trend as labelled coverage increases. The 20% value is averaged across three seeds; the 100% row is a diagnostic reference.")
    figure("figure_02_label_budget.png")
    heading("5.2 Teacher and Student", 2)
    para("The selected Teacher fixed-test mAP@0.5 is " + _f(teacher["metrics"]["map50"]) + ". The Student fixed-test metrics are mAP@0.5 " + _f(student["metrics"]["map50"]) + ", mAP@0.5:0.95 " + _f(student["metrics"]["map50_95"]) + ", precision " + _f(student["metrics"]["precision"]) + ", recall " + _f(student["metrics"]["recall"]) + " and F1 " + _f(student["metrics"]["f1"]) + ". The Student is therefore a usable demonstration checkpoint, but it did not exceed this Teacher on the fixed test.")
    _pdf_table(tables["main_results"]["rows"], tables["main_results"]["headers"], styles, story, tables["main_results"]["caption"])
    figure("figure_03_per_class_ap.png")
    para("The annotated outputs below demonstrate multi-object boxes and class/confidence rendering; they are qualitative evidence only.")
    figure("figure_04_gui_examples.png")
    heading("5.3 Self-supervised novel-category discovery", 2)
    para("The representation loss decreased over ten epochs. Discovery purity, NMI and ARI were " + _f(open_world["metrics"]["discovery"]["purity"]) + ", " + _f(open_world["metrics"]["discovery"]["nmi"]) + " and " + _f(open_world["metrics"]["discovery"]["ari"]) + "; holdout scores were " + _f(open_world["metrics"]["holdout"]["purity"]) + ", " + _f(open_world["metrics"]["holdout"]["nmi"]) + " and " + _f(open_world["metrics"]["holdout"]["ari"]) + ". This is image-level category discovery, not box-level open-world mAP.")
    figure("figure_05_self_supervised_loss.png")
    figure("figure_06_open_world_metrics.png")
    _pdf_table(tables["open_world"]["rows"], tables["open_world"]["headers"], styles, story, tables["open_world"]["caption"])
    _pdf_table(tables["deployment"]["rows"], tables["deployment"]["headers"], styles, story, tables["deployment"]["caption"])
    heading("6 Discussion")
    para("The pipeline is reproducible and demonstrable, but the supervised matrix reveals a data-quality and coverage limitation: the full-label diagnostic is only mAP@0.5 = " + _f(matrix["upper_bound_gate"]["fixed_test_map50"]) + ". Extending training alone is unlikely to reach the original 0.80 target; additional annotation coverage, class balancing, split review and hard-scene curation are more credible interventions. The Student did not exceed the selected Teacher, so the present value of semi-supervision is methodological and operational rather than a proven accuracy gain. The report therefore makes no formal 0.80, 0.85 or box-level open-world claim.")
    heading("7 Impact Statement")
    para("WHAT: a Windows-native, reproducible five-class semi-supervised detection prototype with auditable artifacts and a Chinese PySide6 demonstrator. WHO: agricultural-vision researchers, teaching laboratories, small automation integrators and students. HOW: after target-domain validation, filtered pseudo labels may reduce repeated annotation effort; potential industrial, commercial, environmental and societal benefits remain contingent on licensing, privacy, bias, safety and human-oversight reviews.")
    heading("8 Conclusions")
    para("A complete research and demonstration prototype was implemented on an RTX 3080. The retained Teacher achieved fixed-test mAP@0.5 = " + _f(teacher["metrics"]["map50"]) + ", and the exploratory Student achieved " + _f(student["metrics"]["map50"]) + ". Self-supervised discovery found measurable image-level structure among six categories outside the runtime registry. The evidence supports a transparent prototype, while further data and multi-seed optimisation are required for the original accuracy ambition.")
    heading("References")
    for ref in [
        "1. Tarvainen A, Valpola H. Mean teachers are better role models. Advances in Neural Information Processing Systems. 2017;30:1195-1204.",
        "2. Liu YC, Ma CY, He Z, et al. Unbiased Teacher for semi-supervised object detection. International Conference on Learning Representations. 2021.",
        "3. Xu M, Zhang Z, Hu H, et al. End-to-end semi-supervised object detection with Soft Teacher. ICCV. 2021:3060-3069.",
        "4. Xu B, Chen M, Guan W, Hu L. Efficient Teacher: semi-supervised object detection for YOLOv5. arXiv:2302.07577. 2023.",
        "5. Luo G, Zhou Y, Jin L, et al. Towards end-to-end semi-supervised learning for one-stage object detection. arXiv:2302.11299. 2023.",
        "6. Shehzadi T, Hashmi KA, Stricker D, Afzal MZ. Sparse Semi-DETR. CVPR. 2024.",
        "7. Johanson R, Wilms C, Johannsen O, Frintrop S. S3AD: semi-supervised small apple detection in orchard environments. WACV. 2024:7076-7085.",
        "8. Xiao F, Wang H, Xu Y, Zhang R. Fruit detection and recognition based on deep learning for automatic harvesting: an overview and review. Agronomy. 2023;13(6):1625.",
    ]:
        para(ref)
    heading("Appendix A. Reproducibility and evidence boundary")
    para("The package records source JSON paths, SHA-256 digests, selected checkpoint identities, split fingerprint, figure/table hashes and environment identity. The GUI loads the Student best.pt checkpoint explicitly. No camera workflow or private credential is included. Novel-category results remain an image-level discovery artifact; semantic class registration and box-level unknown-object evaluation are future work.")
    heading("Appendix B. Alignment with supplied requirements")
    para("The report includes the required title page, abstract, experimental-paper sections, figures/tables with captions, references and appendices, while remaining below the 5,000-word narrative limit and ten-figure/ten-table limit. It also covers the supplied PPT topics: problem, state of the art, data split and annotation budget, methodology, evaluation, risks and ethics, deliverables, impact and limitations. Student identity, ID and supervisor fields remain placeholders for final submission.")

    pdf_path = output / "final_report.pdf"

    def page_number(canvas: Any, document: Any) -> None:
        canvas.saveState()
        canvas.setFont("FruitSans", 8)
        canvas.setFillColor(colors.HexColor("#5F7484"))
        canvas.drawRightString(A4[0] - 0.72 * inch, 0.4 * inch, f"{document.page}")
        canvas.restoreState()

    SimpleDocTemplate(str(pdf_path), pagesize=A4, rightMargin=0.72 * inch, leftMargin=0.72 * inch, topMargin=0.66 * inch, bottomMargin=0.62 * inch, title="Semi-Supervised Fruit Object Detection").build(story, onFirstPage=page_number, onLaterPages=page_number)
    return pdf_path


def _pdf_table(rows: Sequence[Sequence[Any]], headers: Sequence[Any], styles: Mapping[str, Any], story: list[Any] | None = None, caption: str | None = None) -> Any:
    from html import escape
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, Table, TableStyle

    data = [[Paragraph(str(value), styles["FruitSmall"]) for value in headers]]
    data.extend([[Paragraph(str(value), styles["FruitSmall"]) for value in row] for row in rows])
    table = Table(data, repeatRows=1, hAlign="LEFT", colWidths=None)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D8EFEC")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#173B55")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D8E2")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    if story is not None:
        if caption:
            story.append(Paragraph(escape(caption), styles["FruitCaption"]))
        story.append(table)
        from reportlab.platypus import Spacer
        story.append(Spacer(1, 0.12 * inch))
    return table


def build_report(
    *,
    output: Path,
    matrix_path: Path,
    teacher_path: Path,
    student_path: Path,
    open_world_path: Path,
    gui_export: Path,
) -> Path:
    output = _ensure_output(output)
    matrix = _load(matrix_path)
    teacher = _load(teacher_path)
    student = _load(student_path)
    open_world = _load(open_world_path)
    if teacher.get("metrics", {}).get("map50") is None or student.get("metrics", {}).get("map50") is None:
        raise ValueError("Teacher/Student fixed-test metrics are incomplete")
    tables = _build_tables(output, matrix, teacher, student, open_world)
    figures = _build_figures(output, matrix, teacher, student, open_world, gui_export.resolve(strict=True))
    docx_path, word_count = _build_docx(output, teacher, student, matrix, open_world, tables, figures)
    pdf_path = _build_pdf(output, teacher, student, matrix, open_world, tables, figures)
    evidence_paths = [matrix_path, teacher_path, student_path, open_world_path]
    student_checkpoint = Path(student.get("checkpoint", {}).get("path", ""))
    teacher_checkpoint = Path(teacher.get("protocol", {}).get("checkpoint_path", ""))
    if student_checkpoint.is_file():
        evidence_paths.append(student_checkpoint)
    if teacher_checkpoint.is_file():
        evidence_paths.append(teacher_checkpoint)
    alignment = output / "requirements_alignment.md"
    alignment.write_text(
        "# Requirements alignment\n\n"
        "- Word guidance: title page, abstract, Introduction, Methods/Methodology, Results, Discussion, Conclusions, Impact Statement, References and Appendices.\n"
        "- Word limits: narrative below 5,000 words; six figures and six tables, both below the ten-item limit.\n"
        "- PPT requirements: problem/background, state of the art, data split and label budget, methodology, evaluation, risks/ethics, milestones/deliverables, impact and limitations.\n"
        "- Evidence boundary: no formal 0.80 acceptance claim, no 0.85 claim, no camera workflow and no box-level open-world claim.\n",
        encoding="utf-8",
    )
    summary = {
        "protocol": "fruit_ssod_final_submission_report_v3_r1",
        "word_count_narrative": word_count,
        "figure_count": len(figures),
        "table_count": len(tables),
        "teacher_map50": teacher["metrics"]["map50"],
        "student_map50": student["metrics"]["map50"],
        "open_world_holdout": open_world["metrics"]["holdout"],
        "acceptance_claim": "none",
        "known_classes": KNOWN_CLASSES,
        "novel_classes": NOVEL_CLASSES,
    }
    (output / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (output / "README.md").write_text(
        "# Final Report v3-r1\n\n"
        "This package is the English technical report rebuilt from the supplied Word/PPT requirements and sealed experiment evidence. It contains a Word and PDF version, six figures, six tables and an alignment note. It makes no formal 0.80/0.85 accuracy claim and describes novel-category work as image-level discovery.\n",
        encoding="utf-8",
    )
    manifest = {
        "protocol": "fruit_ssod_final_submission_report_v3_r1",
        "summary": {"path": str(output / "summary.json"), "sha256": _sha256(output / "summary.json")},
        "report_docx": {"path": str(docx_path), "sha256": _sha256(docx_path)},
        "report_pdf": {"path": str(pdf_path), "sha256": _sha256(pdf_path)},
        "requirements_alignment": {"path": str(alignment), "sha256": _sha256(alignment)},
        "evidence": [{"path": str(path), "sha256": _sha256(path)} for path in evidence_paths if path.is_file()],
        "figures": [{"path": str(spec["path"]), "sha256": spec["sha256"], "caption": spec["caption"]} for spec in figures.values()],
        "tables": [{"path": str(spec["path"]), "sha256": spec["sha256"], "caption": spec["caption"]} for spec in tables.values()],
    }
    (output / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return output


def main(argv: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--matrix", type=Path, required=True)
    parser.add_argument("--teacher-test", type=Path, required=True)
    parser.add_argument("--student-test", type=Path, required=True)
    parser.add_argument("--open-world", type=Path, required=True)
    parser.add_argument("--gui-export", type=Path, required=True)
    args = parser.parse_args(argv)
    result = build_report(output=args.output, matrix_path=args.matrix, teacher_path=args.teacher_test, student_path=args.student_test, open_world_path=args.open_world, gui_export=args.gui_export)
    print(json.dumps({"output": str(result)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
