"""Build a compact DOCX/PDF report from one immutable exploratory package."""
from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path


def _load(package: Path) -> dict:
    return json.loads((package / "summary.json").read_text(encoding="utf-8"))


def _file_evidence(path: Path) -> dict[str, object]:
    payload = path.read_bytes()
    return {
        "path": str(path),
        "bytes": len(payload),
        "sha256": hashlib.sha256(payload).hexdigest(),
    }


def build(package: Path) -> tuple[Path, Path]:
    summary = _load(package)
    runs = summary.get("runs", [])
    selected = summary.get("selected_demo_run_id", "unknown")
    open_world = summary.get("open_world")
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.add_heading("Semi-Supervised Fruit Object Detection", 0)
    doc.add_paragraph("Exploratory research and demonstration report generated from immutable run evidence.")
    doc.add_heading("Scope and limitations", level=1)
    if open_world:
        doc.add_paragraph("The prototype detects Apple, Banana, Orange, Strawberry and Pineapple. It uses offline image inference in a PySide6 desktop GUI and includes a separate post-Student image-level open-world discovery experiment. Camera capture is not part of this delivery. This package is exploratory and does not claim the 0.80 acceptance target.")
    else:
        doc.add_paragraph("The prototype detects Apple, Banana, Orange, Strawberry and Pineapple. It uses offline image inference in a PySide6 desktop GUI; camera and open-world detection are disabled. This package is exploratory and does not claim the 0.80 acceptance target.")
    doc.add_heading("Fixed-test results", level=1)
    table = doc.add_table(rows=1, cols=6)
    for cell, text in zip(table.rows[0].cells, ["Run", "mAP@0.5", "mAP@0.5:0.95", "Precision", "Recall", "F1"]):
        cell.text = text
    for run in runs:
        metrics = run.get("fixed_test", {})
        row = table.add_row().cells
        for cell, key in zip(row, ["run_id", "map50", "map50_95", "precision", "recall", "f1"]):
            value = run.get(key) if key == "run_id" else metrics.get(key)
            cell.text = str(value) if key == "run_id" else f"{float(value):.6f}"
    if open_world:
        results = open_world.get("results", {})
        discovery = results.get("metrics", {}).get("discovery", {})
        holdout = results.get("metrics", {}).get("holdout", {})
        novelty = results.get("novelty", {})
        doc.add_heading("Post-Student open-world discovery", level=1)
        doc.add_paragraph(
            "The independent DeepNIR pool contains six categories outside the five-class runtime registry. "
            "Self-supervised augmentation consistency, deterministic clustering and post-hoc protected-label evaluation "
            "were used to produce reviewable Unknown-cluster proposals. No new runtime class ID was added."
        )
        ow_table = doc.add_table(rows=1, cols=4)
        for cell, text in zip(ow_table.rows[0].cells, ["Split", "Purity", "NMI", "ARI"]):
            cell.text = text
        for label, metrics in (("Discovery", discovery), ("Holdout", holdout)):
            row = ow_table.add_row().cells
            row[0].text = label
            row[1].text = f"{float(metrics.get('purity', 0.0)):.6f}"
            row[2].text = f"{float(metrics.get('nmi', 0.0)):.6f}"
            row[3].text = f"{float(metrics.get('ari', 0.0)):.6f}"
        known_test_count = int(novelty.get("known_fixed_test_count", 0))
        known_fpr_text = (
            f"{float(novelty.get('known_false_positive_rate', 0.0)):.6f} over {known_test_count} known-test images"
            if known_test_count
            else "not measured (no known-test list supplied)"
        )
        doc.add_paragraph(
            f"Novel pool: {results.get('split', {}).get('image_count', 0)} images; "
            f"unknown candidates: {novelty.get('candidate_count', 0)} at threshold {float(novelty.get('threshold', 0.0)):.2f}; "
            f"known-test false-positive rate: {known_fpr_text}."
        )
    doc.add_paragraph(f"Selected demonstration run: {selected}.")
    doc.add_heading("Reproducibility evidence", level=1)
    doc.add_paragraph(f"Protected test split fingerprint: {summary.get('test_split_fingerprint')}.")
    doc.add_paragraph("The package retains the run record, fixed-test JSON, checkpoint SHA-256, GUI export manifest and metadata. Validation/test labels were not used to compose Student training data.")
    docx_path = package / "exploratory_best_result_report.docx"
    doc.save(docx_path)

    # Matplotlib's PDF backend is available in the locked Conda environment and
    # avoids requiring Office/LibreOffice on Windows.
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages
    pdf_path = package / "exploratory_best_result_report.pdf"
    with PdfPages(pdf_path) as pdf:
        fig = plt.figure(figsize=(8.27, 11.69))
        fig.text(0.08, 0.94, "Semi-Supervised Fruit Object Detection", fontsize=18, weight="bold")
        fig.text(0.08, 0.90, "Exploratory research and demonstration report", fontsize=11)
        body = [
            "Five classes: Apple, Banana, Orange, Strawberry, Pineapple.",
            "PySide6 offline image GUI; camera capture is not part of this delivery.",
            "This package is exploratory and does not claim the 0.80 target.",
            "",
            "Fixed-test results:",
        ]
        for run in runs:
            m = run.get("fixed_test", {})
            body.append(f"{run.get('run_id')}: mAP50={float(m.get('map50', 0)):.6f}, Precision={float(m.get('precision', 0)):.6f}, Recall={float(m.get('recall', 0)):.6f}, F1={float(m.get('f1', 0)):.6f}")
        body += ["", f"Selected demo run: {selected}", f"Test split fingerprint: {summary.get('test_split_fingerprint')}", "", "Evidence is bound to the package manifest and checkpoint SHA-256."]
        if open_world:
            results = open_world.get("results", {})
            discovery = results.get("metrics", {}).get("discovery", {})
            holdout = results.get("metrics", {}).get("holdout", {})
            novelty = results.get("novelty", {})
            known_test_count = int(novelty.get("known_fixed_test_count", 0))
            known_fpr_text = (
                f"{float(novelty.get('known_false_positive_rate', 0.0)):.6f} over {known_test_count} known-test images"
                if known_test_count
                else "not measured (no known-test list supplied)"
            )
            body += [
                "",
                "Post-Student open-world discovery (image-level Unknown clusters):",
                f"Novel pool images: {results.get('split', {}).get('image_count', 0)}; categories: {len(results.get('novel_categories_for_protected_evaluation', []))}",
                f"Discovery purity/NMI/ARI: {float(discovery.get('purity', 0.0)):.6f}/{float(discovery.get('nmi', 0.0)):.6f}/{float(discovery.get('ari', 0.0)):.6f}",
                f"Holdout purity/NMI/ARI: {float(holdout.get('purity', 0.0)):.6f}/{float(holdout.get('nmi', 0.0)):.6f}/{float(holdout.get('ari', 0.0)):.6f}",
                f"Unknown candidates: {novelty.get('candidate_count', 0)}; known-test false-positive rate: {known_fpr_text}",
                "Cluster names are post-hoc evaluation mappings; no new runtime class ID was added.",
            ]
        fig.text(0.08, 0.84, "\n".join(body), fontsize=10, va="top", family="DejaVu Sans", wrap=True)
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)
    manifest_path = package / "manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8")) if manifest_path.exists() else {}
    manifest["report_docx"] = _file_evidence(docx_path)
    manifest["report_pdf"] = _file_evidence(pdf_path)
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return docx_path, pdf_path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--package", required=True)
    args = parser.parse_args()
    paths = build(Path(args.package).resolve(strict=True))
    print(json.dumps({"docx": str(paths[0]), "pdf": str(paths[1])}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
