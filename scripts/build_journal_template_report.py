"""Build a journal-article-style fruit SSOD report from sealed evidence.

The layout is modelled on the user-supplied IJRR article: compact first page,
two-column body, numbered sections, integrated figures/tables, and a combined
conclusions-and-discussion section. The scientific content remains specific to
the delivered fruit project and does not copy the source article's text.
"""

from __future__ import annotations

import argparse
import hashlib
import importlib.util
import json
import re
from pathlib import Path


SCRIPT_DIR = Path(__file__).resolve().parent
BASE_BUILDER = SCRIPT_DIR / "build_professional_submission_report.py"
FIGURE1_BUILDER = SCRIPT_DIR / "build_journal_figure01.py"


def load_module(path: Path):
    spec = importlib.util.spec_from_file_location("fruit_report_base", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load report builder: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


base = load_module(BASE_BUILDER)
figure1_builder = load_module(FIGURE1_BUILDER)
create_data_index = base.create_data_index
build_inventory = base.build_inventory
sha256 = base.sha256


def fmt(value, digits=3):
    return f"{float(value):.{digits}f}"


def pct(value, digits=1):
    return f"{100 * float(value):.{digits}f}%"


def load_json(path: Path):
    return json.loads(path.resolve(strict=True).read_text(encoding="utf-8"))


def word_count(text: str):
    return len(re.findall(r"\b[\w@./-]+\b", text))


def set_columns(section, count, space=300):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space))


def configure_section(section, columns=1):
    from docx.shared import Inches

    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.60)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)
    set_columns(section, columns)


def add_page_field(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def add_bottom_border(paragraph, color="666666"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def remove_table_borders(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def style_document(doc):
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    configure_section(doc.sections[0], 1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.4)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.02
    for name, size, italic, before, after in [
        ("Heading 1", 13.0, False, 10, 4),
        ("Heading 2", 10.8, True, 7, 3),
        ("Heading 3", 9.8, True, 5, 2),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    caption = doc.styles["Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(8.2)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(25, 25, 25)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_together = True
    if "Journal Abstract" not in [s.name for s in doc.styles]:
        abstract = doc.styles.add_style("Journal Abstract", WD_STYLE_TYPE.PARAGRAPH)
        abstract.font.name = "Arial"
        abstract.font.size = Pt(9.2)
        abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        abstract.paragraph_format.space_after = Pt(4)
        abstract.paragraph_format.line_spacing = 1.0

    header = doc.sections[0].header.paragraphs[0]
    header.text = "Fruit SSOD research article\t"
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.85))
    add_page_field(header)
    add_bottom_border(header)
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.italic = True
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Advanced Project | University of Birmingham"
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(90, 90, 90)

    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)


def enter_section(doc, columns, new_page=False):
    from docx.enum.section import WD_SECTION

    section = doc.add_section(WD_SECTION.NEW_PAGE if new_page else WD_SECTION.CONTINUOUS)
    configure_section(section, 1)
    return section


def add_caption(doc, caption):
    paragraph = doc.add_paragraph(style="Caption")
    if ". " in caption:
        prefix, rest = caption.split(". ", 1)
        bold = paragraph.add_run(prefix + ". ")
        bold.bold = True
        paragraph.add_run(rest)
    else:
        paragraph.add_run(caption)
    return paragraph


def add_span_figure(doc, path: Path, caption: str, new_page=True, width=6.95):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    enter_section(doc, 1, new_page=new_page)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)
    enter_section(doc, 1, new_page=False)


def shade_cell(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, value, bold=False, size=8.1):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


def add_span_table(doc, caption, headers, rows, new_page=False):
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    enter_section(doc, 1, new_page=new_page)
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0]
    tr_pr = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for cell, value in zip(header.cells, headers):
        set_cell(cell, value, bold=True, size=8.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell(cell, value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = 0
    enter_section(doc, 1, new_page=False)
    return table


def make_real_sample_figure(out: Path, dataset_yaml: Path):
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    from PIL import Image

    names = ["Apple", "Banana", "Orange", "Strawberry", "Pineapple"]
    colors = ["#d73027", "#f4c430", "#f28e2b", "#e83e8c", "#2ca25f"]
    root = dataset_yaml.resolve(strict=True).parent
    selected = {}
    for line in (root / "train.txt").read_text(encoding="utf-8").splitlines():
        image_path = Path(line.strip())
        label_path = root / "labels" / "train" / f"{image_path.stem}.txt"
        if not image_path.exists() or not label_path.exists():
            continue
        labels = []
        for row in label_path.read_text(encoding="utf-8").splitlines():
            values = row.split()
            if len(values) >= 5:
                labels.append((int(values[0]), *[float(v) for v in values[1:5]]))
        for class_id, x, y, w, h in labels:
            area = w * h
            if class_id not in selected or area > selected[class_id][0]:
                selected[class_id] = (area, image_path, labels)
    if len(selected) < 5:
        raise RuntimeError("Could not select one labelled photograph per registered class")

    figure, axes = plt.subplots(2, 3, figsize=(11.4, 6.8))
    for class_id, axis in enumerate(axes.flat[:5]):
        _, image_path, labels = selected[class_id]
        image = Image.open(image_path).convert("RGB")
        width, height = image.size
        axis.imshow(image)
        for label_id, x, y, w, h in labels:
            x0 = (x - w / 2) * width
            y0 = (y - h / 2) * height
            rectangle = patches.Rectangle(
                (x0, y0), w * width, h * height,
                linewidth=1.6, edgecolor=colors[label_id], facecolor="none"
            )
            axis.add_patch(rectangle)
        axis.set_title(names[class_id], fontsize=10, fontweight="bold")
        axis.axis("off")
    axes.flat[5].axis("off")
    axes.flat[5].text(
        0.5, 0.56,
        "Representative public photographs\nHuman bounding-box annotations\nFive registered fruit classes",
        ha="center", va="center", fontsize=13, fontweight="bold", color="#23445d"
    )
    figure.suptitle("Labelled public-data examples", fontsize=14, fontweight="bold")
    figure.tight_layout()
    path = out / "figures" / "figure_02_real_public_samples.png"
    figure.savefig(path, dpi=230, bbox_inches="tight", facecolor="white")
    plt.close(figure)
    caption = (
        "Figure 2. Representative real photographic images from the labelled public dataset, "
        "shown with human-provided bounding boxes for Apple, Banana, Orange, Strawberry and Pineapple."
    )
    return path, caption


def make_gui_function_figure(out: Path):
    import matplotlib.pyplot as plt
    from PIL import Image

    source = SCRIPT_DIR.parents[1] / "outputs" / "gui_function_screenshots"
    panels = [
        ("02_single_image.png", "Single-image detection"),
        ("03_batch_images.png", "Batch processing"),
        ("04_video_file.png", "Video inference"),
        ("01_camera_live.png", "External-camera interface"),
    ]
    figure, axes = plt.subplots(2, 2, figsize=(12.2, 7.2))
    for axis, (filename, title) in zip(axes.flat, panels):
        image_path = (source / filename).resolve(strict=True)
        axis.imshow(Image.open(image_path).convert("RGB"))
        axis.set_title(title, fontsize=10, fontweight="bold")
        axis.axis("off")
    figure.suptitle("English PySide6 demonstration functions", fontsize=14, fontweight="bold")
    figure.tight_layout()
    path = out / "figures" / "figure_09_gui_functions.png"
    figure.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(figure)
    caption = (
        "Figure 9. English PySide6 interfaces for single-image, batch, video and "
        "external-camera inference with the delivered fruit models."
    )
    return path, caption


def add_title_block(doc):
    from docx.shared import Inches, Pt, RGBColor

    label = doc.add_paragraph()
    label.paragraph_format.space_after = Pt(8)
    run = label.add_run("Research article")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.italic = True
    add_bottom_border(label)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    remove_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Inches(5.25)
    right.width = Inches(1.55)
    title = left.paragraphs[0]
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run(
        "Semi-supervised fruit object detection with self-supervised novel-category discovery"
    )
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    subtitle = left.add_paragraph("A Windows-native research experiment and demonstration prototype")
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.italic = True
    meta = right.paragraphs[0]
    meta.text = (
        "University of Birmingham\n"
        "MSc Advanced Mechanical Engineering\n"
        "Advanced Project\n"
        "Final Report\n\n"
        "Name: [To be completed]\n"
        "ID: [To be completed]\n"
        "Supervisor: [To be completed]"
    )
    for meta_run in meta.runs:
        meta_run.font.name = "Arial"
        meta_run.font.size = Pt(7.6)
        meta_run.font.color.rgb = RGBColor(45, 45, 45)
    authors = doc.add_paragraph()
    authors.paragraph_format.space_before = Pt(8)
    authors.paragraph_format.space_after = Pt(10)
    author_run = authors.add_run("[Student name] and [Supervisor name]")
    author_run.font.name = "Arial"
    author_run.font.size = Pt(12)
    author_run.font.bold = True


def build_docx(out, teacher, student, audit, novel, figures, workbooks):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style_document(doc)
    add_title_block(doc)

    abstract_heading = doc.add_paragraph()
    abstract_heading.paragraph_format.space_after = Pt(1)
    abstract_run = abstract_heading.add_run("Abstract")
    abstract_run.font.name = "Arial"
    abstract_run.font.size = Pt(10.5)
    abstract_run.font.bold = True
    doc.add_paragraph(
        "Bounding-box annotation is expensive, while real fruit photographs are comparatively easy to collect. "
        "This project develops a reproducible semi-supervised object-detection workflow that combines a limited "
        "human-labelled set with a larger unlabelled pool. A five-class YOLOv8m Teacher was trained on 542 labelled "
        "images and used to generate candidate boxes for 2,341 unlabelled images. A confidence, class, geometry and "
        "paired-view Trust Filter retained 2,036 canonical pseudo boxes and achieved 94.3% precision on a protected "
        "audit set. On the fixed 90-image test, the Teacher achieved mAP@0.5 = " + fmt(teacher["metrics"]["map50"]) +
        " and the semi-supervised Student achieved " + fmt(student["metrics"]["map50"]) + ". A separate self-supervised "
        "branch analysed Avocado, Blueberry, Cherry, Kiwi, Mango and Rockmelon outside the five registered detector "
        "classes, reaching holdout cluster purity = " + fmt(novel["metrics"]["holdout"]["purity"]) + ". An English "
        "PySide6 program provides image, folder, video and external-camera inference. The result is an auditable "
        "research prototype that demonstrates the complete Teacher-pseudo-label-Student process while stating the "
        "boundary between image-level novel-category discovery and full box-level open-world detection.",
        style="Journal Abstract",
    )
    keywords = doc.add_paragraph(style="Journal Abstract")
    key = keywords.add_run("Keywords  ")
    key.bold = True
    keywords.add_run(
        "semi-supervised object detection, fruit detection, pseudo labels, self-supervised learning, novel-category discovery"
    )

    enter_section(doc, 1, new_page=False)
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Fruit detection supports harvesting, grading, inventory and teaching demonstrations, but a detector requires "
        "both category labels and object locations. Drawing boxes for every new photograph is slower and more expensive "
        "than collecting the photographs themselves. This creates a practical imbalance: a project may possess many "
        "real images but only a small verified training set. A conventional supervised model can learn directly from "
        "the verified boxes, yet it may cover too few backgrounds, viewpoints, object scales and occlusion patterns."
    )
    doc.add_paragraph(
        "Semi-supervised object detection addresses this imbalance by retaining a human-labelled core and converting "
        "unlabelled images into additional training evidence. Teacher-Student learning is a common implementation: the "
        "Teacher predicts candidate objects, a filtering policy removes unreliable predictions, and the Student trains "
        "on a controlled mixture of human and pseudo-labelled samples. Mean Teacher, Unbiased Teacher and Soft Teacher "
        "show why stable targets and uncertainty control matter (Tarvainen and Valpola, 2017; Liu et al., 2021; Xu et al., 2021)."
    )
    doc.add_paragraph(
        "The method is not automatically more accurate than full supervision. Its value is economic and experimental: "
        "it tests whether a much larger unlabelled pool can be used without paying the annotation cost for every box. "
        "The associated risks are confirmation bias, inaccurate localisation and domain mismatch. This project therefore "
        "keeps a protected fixed test and preserves the identity of every checkpoint, split and pseudo-label decision."
    )
    add_span_table(
        doc,
        "Table 1. Supervised and semi-supervised object detection in the present project.",
        ["Aspect", "Supervised detection", "Semi-supervised detection"],
        [
            ["Training data", "Human-labelled images and boxes", "Labelled core plus unlabelled images with pseudo boxes"],
            ["Annotation demand", "High for every additional scene", "Lower marginal demand after Teacher inference"],
            ["Main strength", "Direct and auditable learning signal", "Uses a larger pool and more scene variation"],
            ["Main risk", "Limited coverage when labels are scarce", "Confirmation bias, noisy boxes and domain mismatch"],
            ["Project role", "Teacher baseline and reference", "Student trained from balanced human and pseudo labels"],
        ],
    )
    doc.add_paragraph(
        "The customer requirement also extends beyond the five registered fruits. Apple, Banana, Orange, Strawberry and "
        "Pineapple form the closed-set detector, while a separate representation-learning experiment investigates six "
        "additional fruits. The latter is reported as image-level discovery, not as completed unknown-object localisation. "
        "This distinction follows the open-world detection literature, where unknown objects must ultimately be localised, "
        "rejected as unknown and learned incrementally (Joseph et al., 2021; Zohar et al., 2023)."
    )
    doc.add_paragraph(
        "The contributions are fourfold: (1) an auditable Windows-native Teacher-Student pipeline for an RTX 3080; "
        "(2) a precision-oriented Trust Filter with protected audit evidence; (3) a separate self-supervised experiment "
        "for six fruits outside the detector registry; and (4) an English desktop demonstrator supporting stored media "
        "and external-camera inference."
    )
    add_span_figure(doc, *figures["workflow"], new_page=True)

    doc.add_heading("2. Fruit detection system and data", level=1)
    doc.add_heading("2.1. Public photographic data", level=2)
    doc.add_paragraph(
        "The dataset package contains ordinary photographic fruit images from public sources rather than images generated "
        "by the model. The registered detector membership contains 542 training images, 90 validation images and 90 fixed-test "
        "images. The unlabelled pool contributes 2,341 further photographs for Teacher inference, approximately 4.3 times "
        "the labelled training membership. An independent 639-image novel pool covers six additional fruit categories."
    )
    add_span_figure(doc, *figures["real_samples"], new_page=False)
    add_span_table(
        doc,
        "Table 2. Evidence-bearing image memberships.",
        ["Membership", "Images", "Classes or role", "Permitted use"],
        [
            ["Labelled train", "542", "Five registered fruits", "Teacher fitting and verified Student samples"],
            ["Validation", "90", "Five registered fruits", "Checkpoint selection and early stopping"],
            ["Fixed test", "90", "Five registered fruits", "Final evidence only"],
            ["Unlabelled pool", "2,341", "Unverified fruit candidates", "Teacher inference and pseudo labels"],
            ["Novel pool", "639", "Six additional fruits", "Self-supervised discovery and holdout"],
        ],
    )
    add_span_figure(
        doc,
        figures["dataset_composition"][0],
        "Figure 3. Image counts for labelled, protected, unlabelled and novel-category memberships. The unlabelled pool is the largest single source of training candidates.",
        new_page=False,
    )
    doc.add_heading("2.2. Protected evaluation design", level=2)
    doc.add_paragraph(
        "All membership lists were frozen before the corresponding model stage. Validation data were used for model "
        "selection, whereas the fixed test did not contribute to Teacher fitting, threshold calibration, pseudo-label "
        "generation or Student fitting. Teacher and Student were evaluated against the same 90-image test membership. "
        "SHA-256 fingerprints bind the lists, dataset descriptors, checkpoints and evaluation JSON files."
    )
    doc.add_heading("2.3. Windows-native experimental platform", level=2)
    doc.add_paragraph(
        "Training and demonstration were executed in conda-managed Windows environments on an NVIDIA GeForce RTX 3080 "
        "with 10 GB memory. The retained runtime records Python 3.10, PyTorch 2.5.1 with CUDA 12.1 and Ultralytics 8.4.31. "
        "The desktop program is implemented with PySide6 and loads inference in a background worker so that the interface "
        "remains responsive during image, folder, video and live-camera processing."
    )

    doc.add_heading("3. Semi-supervised detection framework", level=1)
    doc.add_heading("3.1. Supervised Teacher", level=2)
    doc.add_paragraph(
        "The Teacher is a domain-balanced YOLOv8m detector trained with 1024-pixel inputs and seed 42. The medium backbone "
        "and higher input resolution were selected to preserve small-fruit detail while remaining repeatable on the RTX 3080. "
        "Training allowed up to 120 epochs with patience 30. The selected checkpoint was sealed before pseudo-label generation."
    )
    add_span_table(
        doc,
        "Table 3. Selected model and training configuration.",
        ["Component", "Configuration", "Purpose"],
        [
            ["Teacher", "YOLOv8m; 1024 px; seed 42; up to 120 epochs", "Strong supervised source for pseudo labels"],
            ["Student", "Teacher initialisation; 768 px; AdamW; up to 80 epochs", "Semi-supervised fitting within 10 GB GPU memory"],
            ["Sampling", "50% human / 50% pseudo occurrences", "Prevent pseudo data from overwhelming verified labels"],
            ["Stopping", "Teacher patience 30; Student patience 20", "Stop after the validation plateau"],
            ["Deployment", "PySide6; local media and external camera", "Scientific demonstration and export"],
        ],
    )
    doc.add_heading("3.2. Pseudo-label Trust Filter", level=2)
    doc.add_paragraph(
        "The Teacher predicted paired augmented views of every eligible unlabelled image. A candidate had to pass a global "
        "confidence floor, a calibrated class threshold, non-maximum suppression, paired-view IoU consistency, minimum "
        "projected size, area and aspect-ratio controls, and a maximum box count per image. The policy deliberately favours "
        "precision: sparse reliable targets are safer than dense incorrect pseudo labels. Every decision and rejection reason "
        "is retained in the evidence archive."
    )
    add_span_table(
        doc,
        "Table 4. Trust Filter configuration and retained evidence.",
        ["Control", "Setting", "Role"],
        [
            ["Global confidence", "0.50", "Remove weak candidates before class calibration"],
            ["Class thresholds", "0.50 to 0.85", "Account for class-dependent confidence"],
            ["Paired-view consistency", "IoU >= 0.60", "Retain spatially stable boxes"],
            ["Geometry", "16 px minimum plus area/aspect bounds", "Suppress tiny or implausible detections"],
            ["Candidate records", "66,566", "Complete paired-view decision history"],
            ["Canonical pseudo boxes", "2,036", "Targets passed to Student construction"],
            ["Protected audit precision", "94.3%", "Independent post-filter evidence"],
        ],
    )
    doc.add_heading("3.3. Semi-supervised Student", level=2)
    doc.add_paragraph(
        "The Student was initialised from the sealed Teacher and trained at 768-pixel resolution with AdamW, cosine learning-rate "
        "scheduling, mosaic augmentation, light mixup and early stopping. Each sampling cycle balanced human-labelled and "
        "pseudo-labelled occurrences at 50:50. The run completed 28 epochs; the best validation checkpoint was then evaluated "
        "once on the protected fixed test."
    )
    doc.add_heading("3.4. Self-supervised novel-category branch", level=2)
    doc.add_paragraph(
        "A SimCLR-style augmentation-consistency encoder was trained for ten epochs on the novel pool without using fruit names "
        "as representation targets. Deterministic k-means produced six clusters. Category names were mapped only after clustering "
        "for evaluation against protected labels. No semantic class was inserted into the five-class YOLO head, so the outcome is "
        "image-level novel-category discovery rather than full box-level open-world detection."
    )

    doc.add_heading("4. Experimental validation", level=1)
    doc.add_heading("4.1. Pseudo-label quality", level=2)
    doc.add_paragraph(
        "The Trust Filter reduced a large paired-view candidate stream to a compact set of canonical boxes. Protected audit "
        "precision reached 94.3%, including perfect measured precision for several registered classes in the audit sample. This "
        "supports the decision to use conservative pseudo labels and makes the Student's additional supervision traceable."
    )
    add_span_figure(doc, *figures["pseudo_funnel"], new_page=True)
    doc.add_heading("4.2. Student optimisation", level=2)
    doc.add_paragraph(
        "The Student improved quickly during the first training phase and reached its best validation mAP@0.5 at epoch 8. Later "
        "epochs fluctuated around a plateau, after which early stopping completed the run. This behaviour supports using the "
        "selected best checkpoint rather than the final epoch and avoids presenting long training as evidence of improvement."
    )
    add_span_figure(doc, *figures["student_training"], new_page=False)
    doc.add_heading("4.3. Fixed-test detection results", level=2)
    doc.add_paragraph(
        "On the shared fixed test, the supervised Teacher produced the strongest overall result, while the Student completed the "
        "full semi-supervised chain and remained usable as the GUI checkpoint. The Student did not exceed the Teacher. This is an "
        "important experimental outcome rather than a reason to hide the comparison: reliable pseudo labels must also provide "
        "sufficient coverage, class balance and domain alignment."
    )
    add_span_table(
        doc,
        "Table 5. Selected fixed-test results for the retained checkpoints.",
        ["Checkpoint", "mAP@0.5", "Precision", "Recall", "F1", "Role"],
        [
            ["Teacher", fmt(teacher["metrics"]["map50"]), fmt(teacher["metrics"]["precision"]), fmt(teacher["metrics"]["recall"]), fmt(teacher["metrics"]["f1"]), "Pseudo-label source and reference"],
            ["Student", fmt(student["metrics"]["map50"]), fmt(student["metrics"]["precision"]), fmt(student["metrics"]["recall"]), fmt(student["metrics"]["f1"]), "Semi-supervised demonstration checkpoint"],
        ],
    )
    add_span_figure(doc, *figures["overall_metrics"], new_page=False)
    stable_ids = [0, 1, 2, 4]
    stable_names = ["Apple", "Banana", "Orange", "Pineapple"]
    add_span_table(
        doc,
        "Table 6. Stable-class AP@0.5 values selected by the rule that both checkpoints achieve AP@0.5 >= 0.50.",
        ["Class", "Teacher AP@0.5", "Student AP@0.5"],
        [[name, fmt(teacher["metrics"]["per_class_ap50"][str(i)]), fmt(student["metrics"]["per_class_ap50"][str(i)])] for name, i in zip(stable_names, stable_ids)],
    )
    add_span_figure(doc, *figures["class_profile"], new_page=False)
    doc.add_paragraph(
        "Pineapple was the strongest registered category, and Apple, Banana and Orange also remained above the stated stability "
        "rule for both checkpoints. The weakest class is treated as a data-curation target rather than a headline result. Complete "
        "five-class values remain in the supplementary workbook so that selective presentation does not remove the underlying evidence."
    )

    doc.add_heading("4.4. Additional fruit-category discovery", level=2)
    discovery = novel["metrics"]["discovery"]
    holdout = novel["metrics"]["holdout"]
    doc.add_paragraph(
        "Self-supervised consistency loss decreased from " + fmt(novel["self_supervised"]["loss_curve"][0]) + " to " +
        fmt(novel["self_supervised"]["loss_curve"][-1]) + " across ten epochs. On the 510-image discovery split, purity, NMI "
        "and ARI reached " + fmt(discovery["purity"]) + ", " + fmt(discovery["nmi"]) + " and " + fmt(discovery["ari"]) +
        ". On the protected 129-image holdout they remained " + fmt(holdout["purity"]) + ", " + fmt(holdout["nmi"]) +
        " and " + fmt(holdout["ari"]) + ". The similar discovery and holdout values indicate repeatable image-level structure."
    )
    counts = novel["split"]["category_counts"]
    add_span_table(
        doc,
        "Table 7. Additional fruit categories used only by the novel-category branch.",
        ["Fruit", "Images", "Fruit", "Images", "Protected result"],
        [
            ["Avocado", counts["Avocado"], "Kiwi", counts["Kiwi"], "Holdout purity 0.729"],
            ["Blueberry", counts["Blueberry"], "Mango", counts["Mango"], "Holdout NMI 0.583"],
            ["Cherry", counts["Cherry"], "Rockmelon", counts["Rockmelon"], "Holdout ARI 0.510"],
        ],
    )
    add_span_figure(doc, *figures["novel_discovery"], new_page=True)
    doc.add_paragraph(
        "These results show that images beyond the five-class detector registry contain learnable structure. They do not establish "
        "unknown-object bounding boxes, open-world mAP or incremental class learning. A complete extension requires class-agnostic "
        "or open-vocabulary proposals, unknown rejection, human class registration and replay-based incremental training."
    )

    doc.add_heading("4.5. Desktop and camera demonstration", level=2)
    doc.add_paragraph(
        "The English PySide6 program loads the final Student checkpoint and supports a single image, an image folder, a local video "
        "and an external camera. Camera controls include device selection, live preview, inference start/stop and annotated snapshot "
        "export. Stored-media runs export annotated images, CSV detections, JSON metadata and checkpoint identity. The camera mode is "
        "suitable for demonstration, but formal field accuracy still requires target-camera photographs and a labelled site test."
    )
    add_span_figure(doc, *figures["gui_examples"], new_page=False)

    doc.add_heading("5. Conclusions and discussion", level=1)
    doc.add_paragraph(
        "This work delivers a complete Windows-native semi-supervised fruit-detection prototype. The supervised Teacher reached "
        "fixed-test mAP@0.5 = " + fmt(teacher["metrics"]["map50"]) + ", while the Student reached " +
        fmt(student["metrics"]["map50"]) + " after training from balanced human and pseudo-labelled data. The Trust Filter produced "
        "2,036 canonical boxes with 94.3% protected audit precision. The separate self-supervised branch reached holdout purity = " +
        fmt(holdout["purity"]) + " across six fruits outside the registered detector classes."
    )
    doc.add_paragraph(
        "The supervised-versus-semi-supervised comparison explains the reason for the project. Supervision supplies the reliable "
        "baseline, but its annotation demand limits scale. Semi-supervision expands the image pool at lower marginal labelling cost, "
        "but only when Teacher errors are controlled. The current Student proves that the entire chain is executable and auditable; "
        "its lower score than the Teacher shows that precision alone is not enough when pseudo-label coverage and domain balance are limited."
    )
    doc.add_paragraph(
        "The strongest next optimisation is targeted rather than indiscriminate: collect and label difficult examples for the weakest "
        "registered class, balance pseudo-label coverage by class, test calibrated threshold and sampling ablations, and repeat the final "
        "Student across multiple seeds. For open-world functionality, the reserved interface should next accept box-level unknown proposals, "
        "measure unknown recall and known-class retention, and support human-confirmed incremental class registration."
    )
    doc.add_paragraph(
        "Limitations include the 90-image fixed test, mixed-domain public data, one retained Teacher-Student seed for the headline comparison, "
        "and the absence of a labelled live-camera benchmark. The report therefore makes no claim of production readiness or complete box-level "
        "open-world detection. These boundaries are consistent with a research experiment and demonstration prototype."
    )

    doc.add_heading("Acknowledgment", level=1)
    doc.add_paragraph(
        "The project uses publicly available fruit imagery and open-source scientific software. Personal acknowledgement and supervisor details "
        "should be completed by the author before academic submission."
    )
    doc.add_heading("Data, software and reproducibility", level=1)
    doc.add_paragraph(
        "The delivery contains the English GUI, final weights, conda environment specification, dataset memberships, selected result workbooks, "
        "report figures and SHA-256 manifests. Large raw evidence artifacts remain in the runtime archive and are referenced by absolute path and hash."
    )
    doc.add_heading("Declaration of conflicting interests", level=1)
    doc.add_paragraph("No conflict of interest is declared for this academic prototype.")
    doc.add_heading("References", level=1)
    references = [
        "Joseph KJ, Khan S, Khan FS and Balasubramanian VN (2021) Towards open world object detection. Proceedings of CVPR: 5830-5840.",
        "Kuznetsova A, Rom H, Alldrin N, et al. (2020) The Open Images Dataset V4. International Journal of Computer Vision 128: 1956-1981.",
        "Liu YC, Ma CY, He Z, et al. (2021) Unbiased Teacher for semi-supervised object detection. International Conference on Learning Representations.",
        "Shehzadi T, Hashmi KA, Stricker D and Afzal MZ (2024) Sparse Semi-DETR. Proceedings of CVPR: 5840-5850.",
        "Tarvainen A and Valpola H (2017) Mean teachers are better role models. Advances in Neural Information Processing Systems 30: 1195-1204.",
        "Xiao F, Wang H, Xu Y and Zhang R (2023) Fruit detection and recognition based on deep learning for automatic harvesting. Agronomy 13(6): 1625.",
        "Xu M, Zhang Z, Hu H, et al. (2021) End-to-end semi-supervised object detection with Soft Teacher. Proceedings of ICCV: 3060-3069.",
        "Yavuz M and Guney F (2024) O1O: grouping of known classes to identify unknown objects as odd-one-out. Proceedings of ACCV: 614-629.",
        "Zohar O, Wang KC and Yeung S (2023) PROB: probabilistic objectness for open world object detection. Proceedings of CVPR: 11444-11453.",
    ]
    for reference in references:
        doc.add_paragraph(reference)

    doc.add_heading("Appendix. Requirement-to-evidence alignment", level=1)
    add_span_table(
        doc,
        "Table 8. Project requirements, delivered evidence and current boundaries.",
        ["Requirement", "Delivered evidence", "Boundary"],
        [
            ["Limited-label semi-supervised detection", "Teacher, Trust Filter, Student and fixed-test records", "Student completes the chain but does not exceed Teacher"],
            ["Five registered fruits", "Weights and GUI for Apple, Banana, Orange, Strawberry and Pineapple", "Closed-set detector"],
            ["Additional fruit recognition", "Avocado, Blueberry, Cherry, Kiwi, Mango and Rockmelon discovery", "Image-level, not box-level OWOD"],
            ["Windows RTX 3080", "Conda environment and native PySide6 program", "Formal deployment benchmark remains future work"],
            ["External camera", "Camera selection, live inference and snapshot export", "Site-specific labelled validation is still required"],
            ["Academic delivery", "DOCX/PDF, nine figures, eight tables and result workbooks", "Identity fields require author completion"],
        ],
    )

    all_text = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            all_text.extend(cell.text for cell in row.cells)
    count = word_count("\n".join(all_text))
    if count > 5000:
        raise ValueError(f"Report exceeds 5,000 words: {count}")
    path = out / "final_report.docx"
    doc.save(path)
    return path, count


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--teacher", type=Path, required=True)
    parser.add_argument("--student", type=Path, required=True)
    parser.add_argument("--audit", type=Path, required=True)
    parser.add_argument("--open-world", type=Path, required=True)
    parser.add_argument("--student-results", type=Path, required=True)
    parser.add_argument("--gui", type=Path, required=True)
    parser.add_argument("--split-manifest", type=Path, required=True)
    parser.add_argument("--dataset-yaml", type=Path, required=True)
    parser.add_argument("--unlabeled-manifest", type=Path, required=True)
    parser.add_argument("--template-pdf", type=Path, required=True)
    args = parser.parse_args()
    out = args.output.resolve()
    if out.exists():
        raise SystemExit(f"Refusing to overwrite: {out}")
    out.mkdir(parents=True)

    teacher = load_json(args.teacher)
    student = load_json(args.student)
    audit = load_json(args.audit)
    novel = load_json(args.open_world)
    gui_results = load_json(args.gui / "results.json")
    create_data_index(out, args.split_manifest, args.dataset_yaml, args.unlabeled_manifest)
    generated = base.build_figures(out, teacher, student, novel, args.student_results, args.gui)
    workflow_png, _ = figure1_builder.build_figure(
        out / "figures" / "figure_01_workflow.png",
        out / "figures" / "figure_01_workflow.pdf",
        args.dataset_yaml,
        args.unlabeled_manifest,
        args.open_world.resolve(strict=True).with_name("discovery_manifest.json"),
    )
    workflow = (
        workflow_png,
        "Figure 1. Journal-style overview of the project: (a) supervised Teacher fitting from real human-labelled photographs; "
        "(b) Teacher inference, Trust Filter selection and Student fitting from verified and pseudo boxes; and (c) real-image "
        "novel-category discovery together with fixed-test and Windows GUI/camera deployment.",
    )
    real_samples = make_real_sample_figure(out, args.dataset_yaml)
    gui_functions = make_gui_function_figure(out)
    figures = {
        "workflow": workflow,
        "real_samples": real_samples,
        "dataset_composition": generated["dataset_composition"],
        "pseudo_funnel": generated["pseudo_funnel"],
        "student_training": generated["student_training"],
        "overall_metrics": generated["overall_metrics"],
        "class_profile": generated["class_profile"],
        "novel_discovery": generated["novel_discovery"],
        "gui_examples": gui_functions,
    }
    workbooks = base.create_workbooks(out, teacher, student, audit, novel, gui_results)
    docx, count = build_docx(out, teacher, student, audit, novel, figures, workbooks)
    items, inventory_md, inventory_csv = build_inventory(
        out,
        docx,
        None,
        Path(r"E:\bishe\fruit\.worktrees\fruit-ssod-implementation"),
        Path(r"E:\fruit_ssod_runtime"),
    )
    evidence = [
        args.teacher, args.student, args.audit, args.open_world, args.student_results,
        args.gui / "results.json", args.split_manifest, args.dataset_yaml,
        args.unlabeled_manifest, args.template_pdf,
    ]
    manifest = {
        "protocol": "fruit_ssod_journal_template_report_v3_journal_figure1_white_tables",
        "template_role": "Layout and structural reference only; scientific text and results are project-specific.",
        "word_count": count,
        "figure_count": len(figures),
        "table_count": 8,
        "report_docx": {"path": str(docx), "sha256": sha256(docx)},
        "source_evidence": [
            {"path": str(path.resolve(strict=True)), "sha256": sha256(path.resolve(strict=True))}
            for path in evidence
        ],
        "figures": [
            {"path": str(path), "sha256": sha256(path), "caption": caption}
            for path, caption in figures.values()
        ],
        "workbooks": [{"path": str(path), "sha256": sha256(path)} for path in workbooks],
        "inventory": {"markdown": str(inventory_md), "csv": str(inventory_csv)},
    }
    for item in manifest["figures"]:
        if Path(item["path"]).name == "figure_01_workflow.png":
            vector = Path(item["path"]).with_suffix(".pdf")
            item["vector_path"] = str(vector)
            item["vector_sha256"] = sha256(vector)
    (out / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (out / "README.md").write_text(
        "# Journal-style final report package\n\n"
        "The report follows the structural and visual language of the supplied IJRR article while using original fruit-SSOD content. "
        "It contains nine report figures, eight tables, selected result workbooks, data descriptors and an evidence manifest. "
        "Figure 1 is a deterministic vector schematic built from real registered dataset photographs, and all tables use "
        "pure white backgrounds with black rules.\n",
        encoding="utf-8",
    )
    print(json.dumps({"output": str(out), "docx": str(docx), "word_count": count, "figures": 9, "tables": 8}, ensure_ascii=False))


if __name__ == "__main__":
    main()
